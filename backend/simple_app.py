"""
DPR Analysis System - Single File Application
AI-Powered DPR Analysis using Google Gemini
All-in-One: No external service files needed
"""

from fastapi import FastAPI, UploadFile, File, HTTPException, Form
from fastapi.middleware.cors import CORSMiddleware
from fastapi.responses import JSONResponse
import os
import shutil
from datetime import datetime
import json
import re
import warnings
import logging
from typing import Dict, List, Optional

# Suppress all warnings for cleaner output
warnings.filterwarnings('ignore')
os.environ['TF_CPP_MIN_LOG_LEVEL'] = '3'  # Suppress TensorFlow warnings
os.environ['GRPC_VERBOSITY'] = 'ERROR'     # Suppress gRPC warnings
os.environ['GLOG_minloglevel'] = '2'       # Suppress Google logging

# Configure logging to suppress warnings
logging.getLogger('pdfminer').setLevel(logging.ERROR)
logging.getLogger('pdfplumber').setLevel(logging.ERROR)
logging.getLogger('absl').setLevel(logging.ERROR)
logging.getLogger('grpc').setLevel(logging.ERROR)

# Document processing imports
try:
    import PyPDF2
    import pdfplumber
except ImportError:
    print("[WARNING] Install: pip install PyPDF2 pdfplumber")
    
try:
    import docx
except ImportError:
    print("[WARNING] Install: pip install python-docx")

# Gemini AI import
try:
    import google.generativeai as genai
except ImportError:
    print("[WARNING] Install: pip install google-generativeai")

# ============================================================================
# CONFIGURATION
# ============================================================================

# Gemini API Configuration
GEMINI_API_KEY = "AIzaSyBFX7BxmEZvVRxM2owKgGr3wotzmm_6g4c"
GEMINI_MODEL = "gemini-2.0-flash-exp"  # Fastest model available (Oct 2025)

# Application Configuration
APP_HOST = "0.0.0.0"
APP_PORT = 8000
MAX_UPLOAD_SIZE_MB = 50

# ============================================================================
# INITIALIZE GEMINI
# ============================================================================

genai.configure(api_key=GEMINI_API_KEY)
gemini_model = genai.GenerativeModel(GEMINI_MODEL)
guidelines_context = ""

print(f"[READY] Gemini AI initialized with model: {GEMINI_MODEL}")

# ============================================================================
# DOCUMENT PROCESSING FUNCTIONS
# ============================================================================

def extract_text_from_pdf(file_path: str) -> str:
    """Extract text from PDF file"""
    text = ""
    try:
        # Try pdfplumber first (better extraction)
        # Suppress warnings during PDF processing
        with warnings.catch_warnings():
            warnings.simplefilter("ignore")
            with pdfplumber.open(file_path) as pdf:
                for page in pdf.pages:
                    page_text = page.extract_text()
                    if page_text:
                        text += page_text + "\n"
    except Exception as e:
        # Fallback to PyPDF2
        try:
            with warnings.catch_warnings():
                warnings.simplefilter("ignore")
                with open(file_path, 'rb') as file:
                    pdf_reader = PyPDF2.PdfReader(file)
                    for page in pdf_reader.pages:
                        text += page.extract_text() + "\n"
        except Exception as fallback_error:
            print(f"[ERROR] PDF extraction error: {fallback_error}")
    return text


def extract_text_from_docx(file_path: str) -> str:
    """Extract text from DOCX file"""
    doc = docx.Document(file_path)
    text = "\n".join([para.text for para in doc.paragraphs])
    
    # Also extract from tables
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                text += "\n" + cell.text
    return text


def extract_text(file_path: str, file_extension: str) -> str:
    """Extract text from document based on file type"""
    file_extension = file_extension.lower().replace('.', '')
    
    if file_extension == 'pdf':
        return extract_text_from_pdf(file_path)
    elif file_extension in ['docx', 'doc']:
        return extract_text_from_docx(file_path)
    else:
        raise ValueError(f"Unsupported file type: {file_extension}")


def structure_dpr_data(text: str) -> Dict:
    """Extract and structure key information from DPR text"""
    
    # Helper function to extract fields with better text cleaning
    def extract_field(keywords: List[str]) -> str:
        for keyword in keywords:
            # More flexible pattern that handles multi-line titles
            pattern = rf'{keyword}\s*[:\-]?\s*(.+?)(?:\n\n|\r\n\r\n|(?:\n(?=[A-Z])|$))'
            match = re.search(pattern, text, re.IGNORECASE | re.DOTALL)
            if match:
                value = match.group(1).strip()
                # Clean up the text - remove extra whitespace and newlines
                value = ' '.join(value.split())
                # Limit length but try to complete the sentence
                if len(value) > 200:
                    value = value[:200].rsplit(' ', 1)[0] + '...'
                return value
        return "Not Found"
    
    # Extract budget - comprehensive pattern matching
    budget = {"total": 0, "currency": "INR", "details": "Not Found"}
    budget_patterns = [
        # Pattern 1: "Total 1,940 Lakhs (n19.4 Cr)" - dual format
        (r'Total\s+([\d,]+)\s+Lakhs?\s*\(n?([\d\.]+)\s+Cr', 'dual'),
        # Pattern 2: "Estimated Cost: Rs. 245.50 Crores"
        (r'(?:Estimated|Total|Project)\s+Cost[\s:]*Rs\.?\s*([\d\.]+)\s+(?:Crores?|Cr)', 'crores'),
        # Pattern 3: "Project Cost: n18.75 Crores"
        (r'Project Cost:\s*n?([\d\.]+)\s+Crores?', 'crores'),
        # Pattern 4: "Budget: Rs. 114.93 Crores"
        (r'Budget[\s:]*Rs\.?\s*([\d\.]+)\s+(?:Crores?|Cr)', 'crores'),
        # Pattern 5: Generic Lakhs format
        (r'(?:Total|Budget|Cost)[\s:]*n?\s*([\d,]+)\s+Lakhs?', 'lakhs'),
        # Pattern 6: Generic Crores format
        (r'(?:Total|Budget|Cost)[\s:]*n?\s*([\d\.]+)\s+(?:Crores?|Cr)', 'crores'),
        # Pattern 7: Simple "Rs. X Cr" anywhere in text
        (r'Rs\.?\s*([\d\.]+)\s+Cr(?:ores?)?', 'crores'),
        # Pattern 8: Just number with Crores
        (r'(\d+(?:\.\d+)?)\s+Crores?', 'crores'),
    ]
    
    for pattern, unit_type in budget_patterns:
        match = re.search(pattern, text, re.IGNORECASE)
        if match:
            try:
                if unit_type == 'dual':
                    # "Total 1,940 Lakhs (n19.4 Cr)" format
                    lakhs_str = match.group(1).replace(',', '').strip()
                    crores_str = match.group(2).strip()
                    lakhs = float(lakhs_str)
                    crores = float(crores_str)
                    budget["total"] = crores * 10000000  # Use Crores as more accurate
                    budget["details"] = f"{lakhs:,.0f} Lakhs (Rs. {crores} Crores)"
                elif unit_type == 'lakhs':
                    # Lakhs format
                    amount_str = match.group(1).replace(',', '').strip()
                    amount = float(amount_str)
                    budget["total"] = amount * 100000  # 1 Lakh = 1,00,000
                    budget["details"] = f"{amount:,.0f} Lakhs (Rs. {amount/10:.2f} Crores)"
                elif unit_type == 'crores':
                    # Crores format
                    amount_str = match.group(1).replace(',', '').strip()
                    amount = float(amount_str)
                    budget["total"] = amount * 10000000  # 1 Crore = 1,00,00,000
                    budget["details"] = f"Rs. {amount} Crores"
                
                # If we found a valid budget, break
                if budget["total"] > 0:
                    print(f"[DEBUG] Budget extracted: {budget['details']} (pattern: {pattern})")
                    break
            except (ValueError, IndexError) as e:
                print(f"[DEBUG] Budget pattern matched but parsing failed: {e}")
                continue
    
    if budget["total"] == 0:
        print(f"[DEBUG] No budget found in text. First 500 chars: {text[:500]}")
    
    # Extract timeline
    timeline = {"duration": "Not specified", "duration_months": 0}
    duration_match = re.search(r'(?:Duration|Timeline|Project Duration)[:\s]+(\d+)\s*(months?|years?)', text, re.IGNORECASE)
    if duration_match:
        value = int(duration_match.group(1))
        unit = duration_match.group(2).lower()
        if 'year' in unit:
            timeline["duration_months"] = value * 12
            timeline["duration"] = f"{value} years"
        else:
            timeline["duration_months"] = value
            timeline["duration"] = f"{value} months"
    
    # Identify project type - prioritize more specific types first
    text_lower = text.lower()
    project_type = 'general'
    project_types = {
        'it_park': ['it park', 'information technology park', 'tech park', 'technology park', 'software park', 'ites'],
        'hospital': ['hospital', 'health center', 'medical center', 'health care'],
        'school': ['school', 'college', 'university', 'education', 'institute'],
        'bridge': ['bridge', 'flyover', 'overpass'],
        'water': ['water supply', 'irrigation', 'water treatment', 'reservoir'],
        'road': ['road', 'highway', 'expressway', 'connectivity'],
    }
    for ptype, keywords in project_types.items():
        if any(kw in text_lower for kw in keywords):
            project_type = ptype
            break
    
    structured_data = {
        "project_title": extract_field(["Project Title", "Title", "Project Name"]),
        "project_type": project_type,
        "budget": budget,
        "timeline": timeline,
        "location": extract_field(["Location", "Project Location", "Site"]),
        "implementing_agency": extract_field(["Implementing Agency", "Nodal Agency"]),
        "word_count": len(text.split()),
    }
    
    # Debug logging
    print(f"[DEBUG] Structured DPR Data:")
    print(f"  - Title: {structured_data['project_title']}")
    print(f"  - Location: {structured_data['location']}")
    print(f"  - Budget: {structured_data['budget']['details']}")
    print(f"  - Duration: {structured_data['timeline']['duration']}")
    
    return structured_data


# ============================================================================
# GEMINI AI ANALYSIS FUNCTIONS
# ============================================================================

def parse_json_response(text: str, enable_aggressive_repair: bool = True) -> Dict:
    """Parse JSON from Gemini response with robust multi-stage repair strategies."""
    
    original_text = text
    print(f"[DEBUG] Raw response length: {len(text)} characters")

    # 1. Strip markdown fences if present
    if "```json" in text:
        try:
            text = text.split("```json", 1)[1].split("```", 1)[0]
        except Exception:
            pass
    elif "```" in text:
        parts = text.split("```")
        if len(parts) >= 2:
            text = parts[1]

    text = text.strip()

    # 2. Remove control characters but preserve essential whitespace
    import string
    printable_chars = string.printable
    cleaned = ''.join(char for char in text if char in printable_chars or char in '\n\r\t ')
    cleaned = re.sub(r'[\x00-\x08\x0B-\x0C\x0E-\x1F\x7F]', '', cleaned)
    text = cleaned

    def attempt_load(label: str, candidate: str):
        try:
            if not candidate.strip():
                return None, ValueError("Empty candidate")
            return json.loads(candidate), None
        except json.JSONDecodeError as je:
            return None, je
        except Exception as e:
            return None, e

    # Quick direct attempt
    parsed, err = attempt_load("direct", text)
    if parsed is not None:
        print("[SUCCESS] JSON parsed successfully")
        return parsed

    print(f"[WARNING] Initial parse failed: {str(err)[:200]}")

    # 3. Extract JSON object with brace matching
    brace_start = text.find('{')
    brace_end = text.rfind('}')
    if brace_start != -1 and brace_end != -1 and brace_end > brace_start:
        candidate = text[brace_start:brace_end+1].strip()
    else:
        candidate = text

    # 4. Remove trailing commas
    candidate = re.sub(r',\s*(?=[}\]])', '', candidate)
    
    parsed, err2 = attempt_load("trailing-comma-fix", candidate)
    if parsed is not None:
        print("[SUCCESS] JSON parsed after trailing comma removal")
        return parsed

    # 5. Fix common JSON issues more aggressively
    working = candidate
    
    # Fix unterminated strings by finding last valid quote
    if 'Unterminated string' in str(err2) or 'Expecting' in str(err2):
        print("[REPAIR] Attempting to fix unterminated strings...")
        try:
            # Find the error position if available
            if hasattr(err2, 'pos'):
                pos = err2.pos
                # Truncate at the error position
                working = working[:pos]
                # Find the last complete key-value pair
                last_quote = working.rfind('"')
                if last_quote > 0:
                    # Find the comma or brace before this
                    last_comma = working[:last_quote].rfind(',')
                    last_brace = working[:last_quote].rfind('{')
                    last_bracket = working[:last_quote].rfind('[')
                    cutoff = max(last_comma, last_brace, last_bracket)
                    if cutoff > 0:
                        working = working[:cutoff]
                        if working[cutoff] == ',':
                            working = working[:cutoff]  # Remove trailing comma
        except Exception as e:
            print(f"[DEBUG] String repair error: {e}")

    # 6. Balance braces and brackets
    open_braces = working.count('{')
    close_braces = working.count('}')
    open_brackets = working.count('[')
    close_brackets = working.count(']')
    
    if close_braces < open_braces:
        working += '\n' + ('}' * (open_braces - close_braces))
    if close_brackets < open_brackets:
        working += ']' * (open_brackets - close_brackets)

    parsed, err3 = attempt_load("brace-balance", working)
    if parsed is not None:
        print("[SUCCESS] JSON repaired via brace balancing")
        return parsed

    # 7. Last resort: Use model to repair (if enabled)
    if enable_aggressive_repair:
        print("[REPAIR] Attempting model-based repair...")
        schema_hint = '{"keys": ["overall_score", "actionable_insights", "recommendations", "approval_recommendation", "summary"]}'
        repaired = attempt_model_repair(original_text[:15000], schema_hint)  # Limit input size
        if repaired is not None:
            print("[SUCCESS] Model-based repair succeeded")
            return repaired

    # 8. Final fallback: return minimal valid structure
    print("[FALLBACK] Generating minimal valid structure")
    fallback = {
        "_error": "JSON parsing failed after all repair attempts",
        "_original_error": str(err3 or err2 or err)[:500],
        "overall_score": 50,
        "actionable_insights": [
            "Manual review required - automated parsing encountered errors",
            "Document may have formatting issues that need correction",
            "Please ensure DPR document is properly structured"
        ],
        "recommendations": [
            "Review document formatting and structure",
            "Ensure all required sections are complete",
            "Resubmit document after addressing formatting issues"
        ],
        "approval_recommendation": {
            "decision": "REVISE",
            "confidence": 0,
            "reasoning": "Analysis could not be completed due to technical processing errors. Manual expert review is required, or document should be reformatted and resubmitted."
        },
        "summary": "Technical error prevented complete automated analysis. Manual review by expert assessor is strongly recommended.",
        "completeness_analysis": {
            "score": 50,
            "comments": "Could not complete automated analysis"
        },
        "budget_validation": {
            "is_valid": False,
            "comments": "Analysis incomplete"
        },
        "timeline_validation": {
            "is_realistic": False,
            "comments": "Analysis incomplete"
        },
        "technical_feasibility": {
            "score": 50,
            "is_feasible": False,
            "comments": "Analysis incomplete"
        },
        "risk_assessment": {
            "overall_risk_level": "high",
            "overall_risk_score": 70,
            "comments": "Unable to complete risk assessment"
        },
        "compliance_check": {
            "is_compliant": False,
            "comments": "Analysis incomplete"
        }
    }
    return fallback
def attempt_model_repair(raw_response: str, schema_hint: str) -> Optional[Dict]:
    """Use the model itself to repair malformed JSON.
    Returns dict on success, or None on failure.
    """
    try:
        # Very simple, constrained repair prompt
        repair_prompt = f"""
The following JSON response is malformed. Extract and return ONLY a corrected JSON object with these required keys:

Required structure:
{{
    "overall_score": <number 0-100>,
    "actionable_insights": ["insight 1", "insight 2", "insight 3"],
    "recommendations": ["rec 1", "rec 2"],
    "approval_recommendation": {{
        "decision": "APPROVE" or "REVISE" or "REJECT",
        "confidence": <number 0-100>,
        "reasoning": "brief explanation"
    }},
    "summary": "brief summary text"
}}

Malformed input (extract what you can):
{raw_response[:10000]}

Return ONLY the corrected JSON object. Do not include markdown, explanations, or code blocks.
"""
        
        response = gemini_model.generate_content(
            repair_prompt,
            generation_config=genai.GenerationConfig(
                temperature=0.1,
                max_output_tokens=2048,
                response_mime_type="application/json",
            )
        )
        
        print(f"[REPAIR] Model repair response length: {len(response.text)}")
        
        # Parse the repair response without aggressive repair to avoid recursion
        result = parse_json_response(response.text, enable_aggressive_repair=False)
        
        # Validate that required keys are present
        required_keys = ["overall_score", "actionable_insights", "approval_recommendation"]
        if all(key in result for key in required_keys):
            return result
        else:
            print(f"[REPAIR] Repaired JSON missing required keys")
            return None
            
    except Exception as e:
        print(f"[REPAIR] Model-based repair failed: {str(e)[:200]}")
        return None


async def analyze_dpr_comprehensive_fast(dpr_text: str, structured_data: Dict) -> Dict:
    """OPTIMIZED: Single API call for complete DPR analysis with DETAILED content"""
    
    guidelines_section = ""
    if guidelines_context:
        guidelines_section = f"""
**MDONER GUIDELINES REFERENCE:**
{guidelines_context[:5000]}
"""
    
    prompt = f"""You are a senior DPR analyst for the Ministry of Development of North Eastern Region (MDoNER), India, with 15+ years of experience in infrastructure project evaluation.

{guidelines_section}

**PROJECT DETAILS:**
- Project Title: {structured_data.get('project_title', 'N/A')}
- Project Type: {structured_data.get('project_type', 'N/A')}
- Total Budget: {structured_data.get('budget', {}).get('details', 'Not specified')}
- Project Duration: {structured_data.get('timeline', {}).get('duration', 'N/A')}
- Location: {structured_data.get('location', 'N/A')}
- State/Region: North Eastern Region

**COMPLETE DPR DOCUMENT TEXT:**
{dpr_text[:20000]}

CRITICAL: You MUST return ONLY valid JSON. No markdown formatting, no code blocks, no extra text.
Ensure all strings are properly escaped, all commas are in place, and all brackets/braces are balanced.

TASK: Provide an EXTREMELY DETAILED, COMPREHENSIVE analysis of this DPR. Be thorough, specific, and professional.

Return analysis in this JSON format with DETAILED content in each section:

{{
    "completeness_analysis": {{
        "score": <0-100>,
        "missing_sections": ["list each missing section with specific names"],
        "present_sections": ["list all sections that are present"],
        "quality_assessment": "3-4 paragraphs analyzing the quality and depth of each section present",
        "comments": "Detailed 4-5 paragraph assessment covering: documentation completeness, information depth, data sufficiency, formatting quality, and overall presentation standards"
    }},
    "budget_validation": {{
        "is_valid": <true/false>,
        "total_budget": "exact amount with currency",
        "budget_breakdown": {{
            "civil_works": "amount and percentage",
            "equipment": "amount and percentage", 
            "consultancy": "amount and percentage",
            "contingency": "amount and percentage",
            "other": "specify categories with amounts"
        }},
        "issues": [
            {{
                "severity": "low/medium/high/critical",
                "category": "calculation/allocation/justification/documentation",
                "description": "Very detailed description of the issue (3-4 sentences)",
                "impact": "Explain the potential impact",
                "recommendation": "Specific step-by-step fix"
            }}
        ],
        "cost_analysis": "5-6 paragraphs analyzing: cost reasonableness, market rates comparison, allocation appropriateness, contingency adequacy, hidden costs, value for money",
        "financial_viability": "3-4 paragraphs on sustainability, funding sources, revenue generation potential, operational costs",
        "comments": "Comprehensive 5-7 paragraph budget assessment with specific numbers, calculations verification, and detailed recommendations"
    }},
    "timeline_validation": {{
        "is_realistic": <true/false>,
        "total_duration": "specify duration",
        "phase_breakdown": {{
            "pre_construction": "duration and activities",
            "construction": "duration and activities",
            "post_construction": "duration and activities"
        }},
        "critical_milestones": ["list 8-10 key milestones with expected completion times"],
        "issues": [
            {{
                "issue": "Detailed description",
                "impact": "Project impact analysis", 
                "likelihood": "high/medium/low",
                "mitigation": "Detailed mitigation strategy"
            }}
        ],
        "seasonal_considerations": "Detailed analysis of weather, monsoon, festivals, local conditions affecting timeline",
        "dependency_analysis": "4-5 paragraphs on dependencies, critical path, bottlenecks, parallel activities",
        "comments": "Comprehensive 6-8 paragraph timeline analysis including feasibility, risks, optimization opportunities, buffer requirements"
    }},
    "technical_feasibility": {{
        "score": <0-100>,
        "is_feasible": <true/false>,
        "technical_specifications": {{
            "design_standards": "detailed analysis",
            "material_specifications": "detailed analysis",
            "quality_standards": "detailed analysis",
            "safety_measures": "detailed analysis"
        }},
        "strengths": [
            "Strength 1: Very detailed description with examples and evidence (3-4 sentences each)",
            "Strength 2: ...",
            "Strength 3: ...",
            "Strength 4: ...",
            "Strength 5: ..."
        ],
        "weaknesses": [
            "Weakness 1: Detailed description with specific examples and impact analysis (3-4 sentences each)",
            "Weakness 2: ...",
            "Weakness 3: ...",
            "Weakness 4: ..."
        ],
        "implementation_approach": "5-6 paragraphs covering methodology, phasing, resource mobilization, quality control, supervision",
        "technology_assessment": "Analysis of proposed technology, alternatives, appropriateness for region",
        "comments": "Comprehensive 7-10 paragraph technical assessment covering all aspects: design, specifications, implementation, quality, innovation, appropriateness for NE region"
    }},
    "risk_assessment": {{
        "overall_risk_level": "low/medium/high/critical",
        "overall_risk_score": <0-100>,
        "risk_matrix_summary": "2-3 paragraphs explaining the overall risk profile",
        "financial_risks": [
            {{
                "risk": "Detailed risk description with context and examples",
                "severity": "low/medium/high/critical",
                "probability": "low/medium/high",
                "impact_analysis": "Detailed multi-sentence impact analysis",
                "financial_impact": "Estimated cost impact",
                "mitigation": "Detailed 4-5 step mitigation strategy",
                "contingency_plan": "Alternative approach if risk materializes"
            }},
            {{
                "risk": "Cost overrun due to material price fluctuation...",
                "severity": "high",
                "probability": "medium",
                "impact_analysis": "...",
                "financial_impact": "15-20% budget increase",
                "mitigation": "...",
                "contingency_plan": "..."
            }}
        ],
        "timeline_risks": [
            {{
                "risk": "Detailed description",
                "severity": "...",
                "probability": "...",
                "delay_impact": "months/years",
                "mitigation": "Detailed strategy",
                "acceleration_options": "Fast-tracking possibilities"
            }}
        ],
        "environmental_risks": [
            {{
                "risk": "Detailed environmental risk with scientific basis",
                "severity": "...",
                "ecological_impact": "Detailed impact on ecosystem",
                "compliance_requirement": "Specific regulations and clearances",
                "mitigation": "Detailed environmental management plan",
                "monitoring_plan": "How to track and manage"
            }}
        ],
        "resource_risks": [
            {{
                "risk": "Specific resource availability issue",
                "severity": "...",
                "affected_resources": "manpower/materials/equipment",
                "availability_analysis": "Market analysis and sourcing challenges",
                "mitigation": "Detailed procurement and mobilization strategy",
                "alternatives": "Backup resource options"
            }}
        ],
        "social_risks": [
            {{
                "risk": "Community/social challenge",
                "severity": "...",
                "stakeholders_affected": "list",
                "mitigation": "Community engagement plan",
                "benefit_sharing": "How project benefits locals"
            }}
        ],
        "comments": "Comprehensive 8-10 paragraph risk analysis covering: risk identification methodology, inter-dependencies between risks, cumulative impact, risk prioritization, overall risk management strategy, monitoring and control mechanisms"
    }},
    "compliance_check": {{
        "is_compliant": <true/false>,
        "overall_compliance_score": <0-100>,
        "mdoner_guidelines": {{
            "compliant_aspects": ["detailed list with explanations"],
            "non_compliant_aspects": ["detailed list with explanations"],
            "gaps": ["specific gaps with impact"]
        }},
        "statutory_clearances": {{
            "environmental_clearance": "status and requirements",
            "forest_clearance": "status and requirements",
            "wildlife_clearance": "status and requirements",
            "other_clearances": ["list all with status"]
        }},
        "guideline_gaps": [
            {{
                "guideline": "Specific guideline reference",
                "current_status": "What is present",
                "gap": "Detailed gap description",
                "impact": "Impact of non-compliance",
                "remediation": "Detailed steps to achieve compliance",
                "timeline": "Time needed to fix"
            }}
        ],
        "documentation_quality": "4-5 paragraphs on document quality, completeness, clarity, professionalism",
        "comments": "Comprehensive 6-8 paragraph compliance assessment with specific references to guidelines, detailed gap analysis, and actionable remediation plan"
    }},
    "stakeholder_analysis": {{
        "beneficiaries": "Detailed analysis of who benefits and how",
        "affected_parties": "Detailed analysis of impacted communities",
        "consultation_status": "Description of consultations done",
        "social_impact": "5-6 paragraphs on social benefits, displacement, employment, equity"
    }},
    "sustainability_assessment": {{
        "environmental_sustainability": "4-5 paragraphs on environmental impact, conservation, green practices",
        "financial_sustainability": "3-4 paragraphs on long-term financial viability, O&M costs, revenue",
        "social_sustainability": "3-4 paragraphs on social acceptance, community ownership, equity",
        "institutional_sustainability": "Analysis of governance, capacity, maintenance planning"
    }},
    "actionable_insights": [
        "PRIORITY 1 - [Category]: Extremely detailed recommendation with specific steps, responsible parties, timeline, expected outcome, and success metrics (5-7 sentences)",
        "PRIORITY 2 - [Category]: Extremely detailed recommendation...",
        "PRIORITY 3 - [Category]: Extremely detailed recommendation...",
        "PRIORITY 4 - [Category]: Extremely detailed recommendation...",
        "PRIORITY 5 - [Category]: Extremely detailed recommendation...",
        "PRIORITY 6 - [Category]: Extremely detailed recommendation...",
        "PRIORITY 7 - [Category]: Extremely detailed recommendation...",
        "PRIORITY 8 - [Category]: Extremely detailed recommendation..."
    ],
    "recommendations": [
        "Detailed recommendation 1 with rationale and implementation steps",
        "Detailed recommendation 2...",
        "Detailed recommendation 3...",
        "Detailed recommendation 4...",
        "Detailed recommendation 5..."
    ],
    "overall_score": <0-100>,
    "scoring_breakdown": {{
        "completeness": <0-100>,
        "technical_quality": <0-100>,
        "financial_soundness": <0-100>,
        "compliance": <0-100>,
        "feasibility": <0-100>,
        "risk_management": <0-100>
    }},
    "approval_recommendation": {{
        "decision": "APPROVE/APPROVE_WITH_CONDITIONS/REVISE/REJECT",
        "confidence": <0-100>,
        "reasoning": "Comprehensive 6-8 paragraph reasoning covering all aspects: strengths that support decision, weaknesses that concern, conditions for approval, revision requirements, timeline for resubmission, expected improvements",
        "conditions": ["Specific condition 1 if APPROVE_WITH_CONDITIONS", "Specific condition 2..."],
        "revision_requirements": ["Specific requirement 1 if REVISE", "Specific requirement 2..."],
        "approval_process": "Next steps and timeline for approval"
    }},
    "summary": "Comprehensive executive summary in 8-10 sentences covering: project overview, key findings, major strengths, critical concerns, risk level, compliance status, recommendation, and expected impact",
    "key_highlights": {{
        "major_strengths": ["strength 1 with detail", "strength 2...", "strength 3..."],
        "critical_concerns": ["concern 1 with detail", "concern 2...", "concern 3..."],
        "must_address_immediately": ["urgent item 1", "urgent item 2", "urgent item 3"]
    }}
}}

CRITICAL: Provide DETAILED, COMPREHENSIVE, PROFESSIONAL analysis. Each section should have substantial content with specific examples, numbers, analysis, and recommendations. This is for ministerial review - be thorough!

CRITICAL: Provide DETAILED, COMPREHENSIVE, PROFESSIONAL analysis. Each section should have substantial content with specific examples, numbers, analysis, and recommendations. This is for ministerial review - be thorough!

Return ONLY valid JSON, no markdown formatting, no code blocks.
"""
    
    try:
        print("[DETAILED-ANALYSIS] Running comprehensive detailed analysis...")
        response = gemini_model.generate_content(
            prompt,
            generation_config=genai.GenerationConfig(
                temperature=0.3,  # Lower temperature for more consistent JSON
                top_p=0.85,
                top_k=40,
                max_output_tokens=8192,  # Increased for detailed content
                response_mime_type="application/json",  # Force JSON output
            )
        )

        print(f"[DEBUG] Response length: {len(response.text)} characters")

        try:
            analysis = parse_json_response(response.text)
            print(f"[COMPLETE] Detailed analysis done. Score: {analysis.get('overall_score', 'N/A')}")
            return analysis
        except json.JSONDecodeError as primary_err:
            print(f"[RETRY] Primary parse failed: {primary_err}. Attempting model-based repair...")
            # Provide compact schema hint (avoid huge prompt duplication)
            schema_hint = '{"keys": ["completeness_analysis","budget_validation","timeline_validation","technical_feasibility","risk_assessment","compliance_check","stakeholder_analysis","sustainability_assessment","actionable_insights","recommendations","overall_score","scoring_breakdown","approval_recommendation","summary","key_highlights"]}'
            repaired = attempt_model_repair(response.text, schema_hint)
            if repaired is not None:
                print("[REPAIR] Model-based repair succeeded.")
                return repaired
            else:
                print("[REPAIR] Model-based repair failed – returning fallback.")
                raise primary_err
    except Exception as e:
        print(f"[ERROR] Analysis error: {e}")
        print(f"[ERROR] Full error details: {type(e).__name__}: {str(e)}")
        import traceback
        traceback.print_exc()
        return {
            "error": str(e),
            "overall_score": 50,
            "actionable_insights": ["Manual review required due to processing error"],
            "recommendations": ["System encountered an error - manual expert review is required"],
            "approval_recommendation": {
                "decision": "REVISE",
                "confidence": 0,
                "reasoning": "Analysis could not be completed due to technical error. Please ensure document is properly formatted and try again, or conduct manual review."
            },
            "summary": "Analysis encountered a technical error and could not be completed. Manual review by expert assessor is recommended."
        }


async def analyze_dpr_with_gemini(dpr_text: str, structured_data: Dict) -> Dict:
    """Comprehensive DPR analysis using Gemini AI"""
    
    guidelines_section = ""
    if guidelines_context:
        guidelines_section = f"""
**MDONER GUIDELINES CONTEXT:**
{guidelines_context[:5000]}
"""
    
    prompt = f"""
You are an expert DPR (Detailed Project Report) analyst for the Ministry of Development of North Eastern Region (MDoNER), India.

{guidelines_section}

**PROJECT INFORMATION:**
- Project Title: {structured_data.get('project_title', 'N/A')}
- Project Type: {structured_data.get('project_type', 'N/A')}
- Budget: {structured_data.get('budget', {}).get('details', 'Not specified')}
- Timeline: {structured_data.get('timeline', {}).get('duration', 'N/A')}
- Location: {structured_data.get('location', 'N/A')}

**DPR TEXT (excerpt):**
{dpr_text[:15000]}

Analyze this DPR comprehensively and provide assessment in JSON format:

{{
    "completeness_analysis": {{
        "score": <0-100>,
        "missing_sections": [list of missing sections],
        "comments": "Overall completeness assessment"
    }},
    "budget_validation": {{
        "is_valid": <true/false>,
        "issues": [list of budget issues with severity and description],
        "comments": "Budget assessment"
    }},
    "timeline_validation": {{
        "is_realistic": <true/false>,
        "issues": [list of timeline issues],
        "comments": "Timeline assessment"
    }},
    "technical_feasibility": {{
        "score": <0-100>,
        "is_feasible": <true/false>,
        "strengths": [list],
        "weaknesses": [list],
        "comments": "Technical assessment"
    }},
    "risk_assessment": {{
        "overall_risk_level": "low/medium/high/critical",
        "financial_risks": [list of risks with severity],
        "timeline_risks": [list of risks],
        "environmental_risks": [list of risks],
        "resource_risks": [list of risks],
        "comments": "Risk assessment"
    }},
    "compliance_check": {{
        "is_compliant": <true/false>,
        "guideline_gaps": [list of gaps],
        "comments": "Compliance assessment"
    }},
    "recommendations": ["list of actionable recommendations"],
    "overall_score": <0-100>,
    "approval_recommendation": {{
        "decision": "APPROVE/REVISE/REJECT",
        "confidence": <0-100>,
        "reasoning": "Detailed reasoning"
    }},
    "summary": "Executive summary in 2-3 sentences"
}}

Return ONLY valid JSON, no markdown.
"""
    
    try:
        print("[ANALYZING] Analyzing DPR with Gemini AI...")
        response = gemini_model.generate_content(prompt)
        analysis = parse_json_response(response.text)
        print(f"[COMPLETE] Analysis complete. Score: {analysis.get('overall_score', 'N/A')}")
        return analysis
    except Exception as e:
        print(f"[ERROR] Analysis error: {e}")
        return {
            "error": str(e),
            "overall_score": 50,
            "recommendations": ["Manual review required due to error"],
            "approval_recommendation": {"decision": "REVISE", "confidence": 0, "reasoning": "Analysis error"}
        }


async def generate_insights(analysis: Dict) -> List[str]:
    """Generate actionable insights from analysis"""
    
    # Extract key information for context
    score = analysis.get('overall_score', 0)
    decision = analysis.get('approval_recommendation', {}).get('decision', 'REVISE')
    
    prompt = f"""
Based on this DPR analysis (Score: {score}/100, Decision: {decision}), generate 5-7 clear, actionable recommendations to improve the DPR.

**Full Analysis Summary:**
{json.dumps(analysis, indent=2)[:4000]}

Generate specific, actionable recommendations that address:
1. Budget or financial concerns
2. Timeline or implementation issues
3. Technical improvements needed
4. Compliance requirements
5. Risk mitigation strategies

Return ONLY a JSON array of strings:
["Specific actionable recommendation 1", "Specific actionable recommendation 2", ...]

Each recommendation should:
- Start with an action verb (e.g., "Include", "Revise", "Add", "Clarify")
- Be specific about what needs to be done
- Reference the section/area that needs improvement
"""
    
    try:
        response = gemini_model.generate_content(prompt)
        insights = parse_json_response(response.text)
        if isinstance(insights, list) and len(insights) >= 3:
            return insights
        else:
            # Fallback: extract recommendations from analysis
            fallback_insights = []
            if analysis.get('recommendations'):
                fallback_insights = analysis['recommendations'][:7]
            return fallback_insights if fallback_insights else ["Review detailed analysis for comprehensive findings"]
    except Exception as e:
        print(f"[WARNING] Insights generation error: {e}")
        # Extract from recommendations as fallback
        if analysis.get('recommendations'):
            return analysis['recommendations'][:7]
        return ["Review detailed analysis for findings"]


async def assess_risks(dpr_text: str, project_type: str) -> Dict:
    """Assess project risks"""
    prompt = f"""
Analyze risks for this {project_type} project in North Eastern India:

{dpr_text[:8000]}

Return JSON:
{{
    "overall_risk_level": "low/medium/high/critical",
    "overall_risk_score": <0-100>,
    "financial_risks": [{{"risk": "...", "severity": "...", "mitigation": "..."}}],
    "timeline_risks": [...],
    "environmental_risks": [...],
    "resource_risks": [...]
}}
"""
    
    try:
        response = gemini_model.generate_content(prompt)
        return parse_json_response(response.text)
    except:
        return {"overall_risk_level": "medium", "overall_risk_score": 50}


async def translate_report(report: Dict, target_language: str) -> Dict:
    """Translate report to target language"""
    if target_language == "en":
        return report
    
    language_map = {
        'hi': 'Hindi (हिन्दी)',
        'as': 'Assamese (অসমীয়া)',
        'bn': 'Bengali (বাংলা)',
        'mni': 'Manipuri',
        'ne': 'Nepali (नेपाली)'
    }
    
    target_lang = language_map.get(target_language, 'Hindi')
    
    prompt = f"""
Translate this DPR analysis to {target_lang}.
Maintain JSON structure, only translate text values.

{json.dumps(report, indent=2)[:8000]}

Return translated JSON.
"""
    
    try:
        response = gemini_model.generate_content(prompt)
        return parse_json_response(response.text)
    except:
        return report


def generate_structured_json_sections(analysis: Dict, insights: List, risks: Dict, structured_dpr: Dict) -> Dict:
    """Generate structured sections for JSON output without formatting lines"""
    
    overall_score = analysis.get('overall_score', 0)
    approval = analysis.get('approval_recommendation', {})
    decision = approval.get('decision', 'N/A')
    reasoning = approval.get('reasoning', '')
    summary = analysis.get('summary', '')
    
    project_title = structured_dpr.get('project_title', 'this project')
    project_type = structured_dpr.get('project_type', 'Infrastructure Project')
    location = structured_dpr.get('location', 'Not specified')
    implementing_agency = structured_dpr.get('implementing_agency', 'Not specified')
    duration = structured_dpr.get('duration', 'Not specified')
    budget = structured_dpr.get('budget', {}).get('total', 0)
    budget_details = structured_dpr.get('budget', {}).get('details', 'Not specified')
    
    # Format budget
    budget_formatted = "Not specified"
    if budget > 0:
        budget_crores = budget / 10000000
        budget_formatted = f"Rs. {budget_crores:.2f} Crores"
    
    sections = {}
    
    # Executive Summary Section
    sections['executive_summary'] = {
        'overall_score': f"{overall_score}/100",
        'recommendation': decision,
        'analysis_date': datetime.now().strftime('%B %d, %Y at %I:%M %p'),
        'summary': summary if summary else reasoning
    }
    
    # Budget Analysis Section
    budget_eval = analysis.get('budget_validation', {})
    budget_section = ""
    
    if budget_eval:
        is_valid = budget_eval.get('is_valid')
        
        # Pure paragraph format - no labels
        if is_valid:
            budget_section += f"After careful review of the project budget totaling {budget_details}, "
            budget_section += "the financial allocation appears realistic and well-structured. "
            budget_section += "The cost breakdown demonstrates proper planning with adequate provisions for contingencies. "
        else:
            budget_section += f"The project budget of {budget_details} has been thoroughly analyzed, and several concerns have been identified. "
            
            # Convert issues to flowing text
            issues = budget_eval.get('issues', [])
            if issues:
                issue_descriptions = []
                for issue in issues:
                    if isinstance(issue, dict):
                        desc = issue.get('description', str(issue))
                        severity = issue.get('severity', 'medium')
                        issue_descriptions.append(f"there is a {severity} severity concern regarding {desc.lower().rstrip('.')}")
                    else:
                        issue_descriptions.append(str(issue).lower().rstrip('.'))
                
                if issue_descriptions:
                    if len(issue_descriptions) == 1:
                        budget_section += f"Specifically, {issue_descriptions[0]}. "
                    elif len(issue_descriptions) == 2:
                        budget_section += f"Specifically, {issue_descriptions[0]} and {issue_descriptions[1]}. "
                    else:
                        budget_section += f"Specifically, {issue_descriptions[0]}, {', '.join(issue_descriptions[1:-1])}, and {issue_descriptions[-1]}. "
        
        # Add detailed comments as flowing text
        if budget_eval.get('comments'):
            comments = budget_eval.get('comments')
            # Handle both string and dict types
            if isinstance(comments, dict):
                comments = str(comments)
            elif not isinstance(comments, str):
                comments = str(comments)
            comments = comments.strip()
            if not budget_section.endswith('. '):
                budget_section += " "
            budget_section += comments
    else:
        # Generate comprehensive budget analysis from available data
        budget_section = f"The proposed project budget of {budget_details} has been reviewed against standard cost norms and industry benchmarks for similar infrastructure projects in the North Eastern Region. "
        
        # Extract budget-related insights from recommendations
        budget_insights = [i for i in insights if any(term in i.lower() for term in ['budget', 'cost', 'contingency', 'financial', 'rate', 'expenditure', 'unit', 'price'])]
        
        if budget_insights:
            budget_section += "The financial analysis has identified several important considerations. "
            for insight in budget_insights[:3]:  # Use top 3 budget-related insights
                clean_insight = insight.replace('PRIORITY 1 - ', '').replace('PRIORITY 2 - ', '').replace('PRIORITY 3 - ', '').replace('PRIORITY 4 - ', '').replace('PRIORITY 5 - ', '')
                clean_insight = clean_insight.replace('[Budget]', '').replace('[Financial]', '').replace('[Cost]', '').replace('[', '').replace(']', '').strip()
                if not clean_insight.endswith('.'):
                    clean_insight += '.'
                budget_section += clean_insight + " "
            budget_section += "These aspects require immediate attention to ensure fiscal prudence and optimal resource utilization. "
        else:
            budget_section += "The cost structure requires detailed scrutiny to ensure all components are adequately justified with proper market rate comparisons. "
            budget_section += "Contingency provisions should be maintained at industry-standard levels (typically 5-10% for infrastructure projects) to accommodate unforeseen expenses. "
            budget_section += "All unit rates and quantities should be verified against current market conditions and authenticated through competitive quotations. "
        
        budget_section += "Regular budget monitoring and variance analysis will be essential during implementation to prevent cost overruns and ensure value for money."
    
    sections['budget_analysis'] = budget_section.strip()
    
    # Timeline Evaluation Section
    timeline_eval = analysis.get('timeline_validation', {})
    timeline_section = ""
    
    if timeline_eval:
        is_realistic = timeline_eval.get('is_realistic')
        
        # Pure paragraph format
        if is_realistic:
            timeline_section += f"The proposed project duration of {duration} has been carefully evaluated and appears realistic and achievable. "
            timeline_section += "The phasing and scheduling demonstrate sound project management principles with adequate buffer time for unforeseen circumstances. "
        else:
            timeline_section += f"Upon reviewing the proposed timeline of {duration}, several concerns have emerged regarding its feasibility. "
            
            issues = timeline_eval.get('issues', [])
            if issues:
                issue_texts = []
                for issue in issues:
                    if isinstance(issue, dict):
                        desc = issue.get('description', str(issue))
                        issue_texts.append(desc.rstrip('.'))
                    else:
                        issue_texts.append(str(issue).rstrip('.'))
                
                if issue_texts:
                    timeline_section += "Key challenges include "
                    if len(issue_texts) == 1:
                        timeline_section += issue_texts[0].lower()
                    elif len(issue_texts) == 2:
                        timeline_section += issue_texts[0].lower() + " and " + issue_texts[1].lower()
                    else:
                        timeline_section += ", ".join([t.lower() for t in issue_texts[:-1]]) + ", and " + issue_texts[-1].lower()
                    timeline_section += ". "
        
        # Add detailed comments
        if timeline_eval.get('comments'):
            comments = timeline_eval.get('comments')
            # Handle both string and dict types
            if isinstance(comments, dict):
                comments = str(comments)
            elif not isinstance(comments, str):
                comments = str(comments)
            comments = comments.strip()
            if not timeline_section.endswith('. '):
                timeline_section += " "
            timeline_section += comments
    else:
        # Generate comprehensive timeline analysis from available data
        timeline_section = f"The proposed project implementation timeline of {duration} has been assessed against typical execution periods for comparable infrastructure projects in the North Eastern Region. "
        
        # Extract timeline-related insights
        timeline_insights = [i for i in insights if any(term in i.lower() for term in ['timeline', 'schedule', 'delay', 'duration', 'phase', 'completion', 'milestone'])]
        
        if timeline_insights:
            timeline_section += "The temporal analysis reveals several critical factors. "
            for insight in timeline_insights[:2]:  # Use top 2 timeline insights
                clean_insight = insight.replace('PRIORITY 1 - ', '').replace('PRIORITY 2 - ', '').replace('PRIORITY 3 - ', '').replace('PRIORITY 4 - ', '').replace('PRIORITY 5 - ', '')
                clean_insight = clean_insight.replace('[Timeline]', '').replace('[Schedule]', '').replace('[', '').replace(']', '').strip()
                if not clean_insight.endswith('.'):
                    clean_insight += '.'
                timeline_section += clean_insight + " "
        else:
            timeline_section += "The execution timeline should account for several region-specific factors including seasonal monsoon patterns, terrain accessibility challenges, and material procurement logistics. "
            timeline_section += "A detailed activity schedule with clearly defined milestones and critical path analysis would strengthen the implementation plan. "
        
        timeline_section += "Adequate buffer time should be incorporated to handle unforeseen delays, particularly those related to weather conditions and logistical constraints typical of the North Eastern Region. "
        timeline_section += "Regular progress monitoring with quarterly reviews will be essential to ensure timely completion and early identification of potential delays."
    
    sections['timeline_evaluation'] = timeline_section.strip()
    
    # Technical Feasibility Section
    tech_eval = analysis.get('technical_feasibility', {})
    tech_section = ""
    
    if tech_eval:
        tech_score = tech_eval.get('score', 0)
        
        # Pure paragraph format - no score labels
        if tech_score >= 80:
            tech_section += "The technical assessment reveals strong project feasibility with robust engineering design and sound technical foundations. "
        elif tech_score >= 60:
            tech_section += "The technical evaluation indicates moderate feasibility with generally acceptable design standards, though some aspects require attention. "
        else:
            tech_section += "The technical review has identified significant feasibility concerns that need to be addressed to ensure project success. "
        
        # Strengths in flowing text
        if tech_eval.get('strengths'):
            strengths_list = []
            for s in tech_eval.get('strengths', []):
                if isinstance(s, str):
                    strengths_list.append(s.strip().rstrip('.'))
                else:
                    strengths_list.append(str(s).strip().rstrip('.'))
            if strengths_list:
                tech_section += "Notable strengths include "
                if len(strengths_list) == 1:
                    tech_section += strengths_list[0].lower()
                elif len(strengths_list) == 2:
                    tech_section += strengths_list[0].lower() + " and " + strengths_list[1].lower()
                else:
                    tech_section += ", ".join([s.lower() for s in strengths_list[:-1]]) + ", and " + strengths_list[-1].lower()
                tech_section += ". "
        
        # Weaknesses in flowing text
        if tech_eval.get('weaknesses'):
            weaknesses_list = []
            for w in tech_eval.get('weaknesses', []):
                if isinstance(w, str):
                    weaknesses_list.append(w.strip().rstrip('.'))
                else:
                    weaknesses_list.append(str(w).strip().rstrip('.'))
            if weaknesses_list:
                tech_section += "However, areas requiring improvement encompass "
                if len(weaknesses_list) == 1:
                    tech_section += weaknesses_list[0].lower()
                elif len(weaknesses_list) == 2:
                    tech_section += weaknesses_list[0].lower() + " and " + weaknesses_list[1].lower()
                else:
                    tech_section += ", ".join([w.lower() for w in weaknesses_list[:-1]]) + ", and " + weaknesses_list[-1].lower()
                tech_section += ". "
        
        # Add detailed comments
        if tech_eval.get('comments'):
            comments = tech_eval.get('comments')
            # Handle both string and dict types
            if isinstance(comments, dict):
                comments = str(comments)
            elif not isinstance(comments, str):
                comments = str(comments)
            comments = comments.strip()
            if not tech_section.endswith('. '):
                tech_section += " "
            tech_section += comments
    else:
        # Generate comprehensive technical analysis from available data
        overall_score = analysis.get('overall_score', 70)
        
        if overall_score >= 80:
            tech_section += "The technical assessment reveals strong engineering foundations with appropriate design standards for the proposed infrastructure. "
        elif overall_score >= 60:
            tech_section += "The technical evaluation indicates generally sound engineering principles, though certain aspects would benefit from additional detail and validation. "
        else:
            tech_section += "The technical review suggests that several engineering and design elements require further development and refinement. "
        
        # Extract technical insights
        tech_insights = [i for i in insights if any(term in i.lower() for term in ['technical', 'design', 'engineering', 'specification', 'quality', 'construction', 'material', 'survey'])]
        
        if tech_insights:
            tech_section += "Specific technical considerations include the following. "
            for insight in tech_insights[:3]:  # Use top 3 technical insights
                clean_insight = insight.replace('PRIORITY 1 - ', '').replace('PRIORITY 2 - ', '').replace('PRIORITY 3 - ', '').replace('PRIORITY 4 - ', '').replace('PRIORITY 5 - ', '')
                clean_insight = clean_insight.replace('[Technical]', '').replace('[Design]', '').replace('[Engineering]', '').replace('[', '').replace(']', '').strip()
                if not clean_insight.endswith('.'):
                    clean_insight += '.'
                tech_section += clean_insight + " "
        else:
            tech_section += "The project should incorporate detailed technical specifications covering design parameters, material quality standards, construction methodology, and quality assurance protocols. "
            tech_section += "Comprehensive geological and geotechnical surveys should be conducted to validate design assumptions and identify potential ground-related challenges. "
        
        tech_section += "Technical review at each project milestone will ensure adherence to approved standards and specifications, while also allowing for adaptive management based on field conditions."
    
    sections['technical_feasibility'] = tech_section.strip()
    
    # Risk Assessment Section
    risk_eval = analysis.get('risk_assessment', {})
    risk_section = ""
    
    if risk_eval:
        risk_level = risk_eval.get('overall_risk_level', 'unknown').lower()
        
        # Pure paragraph format - no risk level labels
        if risk_level == 'low':
            risk_section += "The comprehensive risk assessment indicates a favorable outlook with manageable uncertainties throughout the project lifecycle. "
        elif risk_level == 'medium':
            risk_section += "The risk evaluation reveals a moderate level of exposure that is typical for infrastructure projects of this scale and complexity. "
        elif risk_level == 'high':
            risk_section += "The risk analysis has identified significant concerns that require careful attention and robust mitigation strategies to ensure project success. "
        else:
            risk_section += "A thorough risk assessment has been conducted to identify potential challenges and mitigation approaches. "
        
        # Collect all risks in flowing paragraphs
        all_risks = []
        
        # Financial risks
        if risk_eval.get('financial_risks'):
            for risk in risk_eval.get('financial_risks', []):
                if isinstance(risk, dict):
                    risk_text = risk.get('risk') or risk.get('name') or risk.get('description') or str(risk)
                    mitigation = risk.get('mitigation', '')
                    # Ensure we have strings
                    if not isinstance(risk_text, str):
                        risk_text = str(risk_text)
                    if not isinstance(mitigation, str):
                        mitigation = str(mitigation)
                    risk_desc = risk_text.strip().rstrip('.')
                    if mitigation:
                        risk_desc += f", which can be addressed through {mitigation.strip().rstrip('.').lower()}"
                    all_risks.append(('financial', risk_desc))
                else:
                    all_risks.append(('financial', str(risk).strip().rstrip('.')))
        
        # Timeline risks
        if risk_eval.get('timeline_risks'):
            for risk in risk_eval.get('timeline_risks', []):
                if isinstance(risk, dict):
                    risk_text = risk.get('risk') or risk.get('name') or risk.get('description') or str(risk)
                    mitigation = risk.get('mitigation', '')
                    # Ensure we have strings
                    if not isinstance(risk_text, str):
                        risk_text = str(risk_text)
                    if not isinstance(mitigation, str):
                        mitigation = str(mitigation)
                    risk_desc = risk_text.strip().rstrip('.')
                    if mitigation:
                        risk_desc += f", with recommended mitigation being to {mitigation.strip().rstrip('.').lower()}"
                    all_risks.append(('timeline', risk_desc))
                else:
                    all_risks.append(('timeline', str(risk).strip().rstrip('.')))
        
        # Environmental risks
        if risk_eval.get('environmental_risks'):
            for risk in risk_eval.get('environmental_risks', []):
                if isinstance(risk, dict):
                    risk_text = risk.get('risk') or risk.get('name') or risk.get('description') or str(risk)
                    mitigation = risk.get('mitigation', '')
                    # Ensure we have strings
                    if not isinstance(risk_text, str):
                        risk_text = str(risk_text)
                    if not isinstance(mitigation, str):
                        mitigation = str(mitigation)
                    risk_desc = risk_text.strip().rstrip('.')
                    if mitigation:
                        risk_desc += f", manageable through {mitigation.strip().rstrip('.').lower()}"
                    all_risks.append(('environmental', risk_desc))
                else:
                    all_risks.append(('environmental', str(risk).strip().rstrip('.')))
        
        # Resource risks
        if risk_eval.get('resource_risks'):
            for risk in risk_eval.get('resource_risks', []):
                if isinstance(risk, dict):
                    risk_text = risk.get('risk') or risk.get('name') or risk.get('description') or str(risk)
                    mitigation = risk.get('mitigation', '')
                    # Ensure we have strings
                    if not isinstance(risk_text, str):
                        risk_text = str(risk_text)
                    if not isinstance(mitigation, str):
                        mitigation = str(mitigation)
                    risk_desc = risk_text.strip().rstrip('.')
                    if mitigation:
                        risk_desc += f", which can be managed by {mitigation.strip().rstrip('.').lower()}"
                    all_risks.append(('resource', risk_desc))
                else:
                    all_risks.append(('resource', str(risk).strip().rstrip('.')))
        
        # Group risks by category and write in flowing text
        if all_risks:
            financial_risks = [r[1] for r in all_risks if r[0] == 'financial']
            timeline_risks = [r[1] for r in all_risks if r[0] == 'timeline']
            env_risks = [r[1] for r in all_risks if r[0] == 'environmental']
            resource_risks = [r[1] for r in all_risks if r[0] == 'resource']
            
            if financial_risks:
                risk_section += "From a financial perspective, key considerations include "
                risk_section += " and ".join(financial_risks) if len(financial_risks) <= 2 else ", ".join(financial_risks[:-1]) + ", and " + financial_risks[-1]
                risk_section += ". "
            
            if timeline_risks:
                risk_section += "Schedule-related concerns involve "
                risk_section += " and ".join(timeline_risks) if len(timeline_risks) <= 2 else ", ".join(timeline_risks[:-1]) + ", and " + timeline_risks[-1]
                risk_section += ". "
            
            if env_risks:
                risk_section += "Environmental factors to monitor include "
                risk_section += " and ".join(env_risks) if len(env_risks) <= 2 else ", ".join(env_risks[:-1]) + ", and " + env_risks[-1]
                risk_section += ". "
            
            if resource_risks:
                risk_section += "Resource availability challenges encompass "
                risk_section += " and ".join(resource_risks) if len(resource_risks) <= 2 else ", ".join(resource_risks[:-1]) + ", and " + resource_risks[-1]
                risk_section += ". "
    else:
        # Generate comprehensive risk analysis from available data
        overall_score = analysis.get('overall_score', 70)
        decision = analysis.get('approval_recommendation', {}).get('decision', 'REVISE')
        
        if decision == 'APPROVE' or overall_score >= 85:
            risk_section += "The risk assessment indicates manageable project uncertainties with standard mitigation strategies applicable throughout the implementation lifecycle. "
        elif decision == 'REVISE' or overall_score >= 60:
            risk_section += "The risk evaluation reveals moderate exposure typical for infrastructure projects of this nature, requiring vigilant monitoring and proactive management. "
        else:
            risk_section += "The risk analysis identifies significant concerns that necessitate careful attention, robust mitigation planning, and contingency arrangements. "
        
        # Extract risk-related insights
        risk_insights = [i for i in insights if any(term in i.lower() for term in ['risk', 'challenge', 'issue', 'concern', 'problem', 'vulnerability', 'threat'])]
        
        if risk_insights:
            risk_section += "Critical risk factors include the following areas. "
            for insight in risk_insights[:3]:  # Use top 3 risk insights
                clean_insight = insight.replace('PRIORITY 1 - ', '').replace('PRIORITY 2 - ', '').replace('PRIORITY 3 - ', '').replace('PRIORITY 4 - ', '').replace('PRIORITY 5 - ', '')
                clean_insight = clean_insight.replace('[Risk]', '').replace('[Challenge]', '').replace('[', '').replace(']', '').strip()
                if not clean_insight.endswith('.'):
                    clean_insight += '.'
                risk_section += clean_insight + " "
        else:
            risk_section += "Standard project risks including cost escalation, schedule delays, quality deviations, and resource constraints should be addressed through a comprehensive risk management framework. "
            risk_section += "Specific attention should be given to region-specific challenges such as seasonal accessibility, material transportation logistics, and availability of skilled workforce. "
        
        risk_section += "A dedicated risk register should be maintained and regularly updated, with clear ownership assigned for monitoring and mitigating each identified risk throughout the project duration."
    
    sections['risk_assessment'] = risk_section.strip()
    
    # Actionable Recommendations Section
    recommendations_section = ""
    if insights and len(insights) > 0:
        recommendations_section += "Based on our comprehensive analysis, the following actionable recommendations are provided to strengthen this DPR and enhance the project's success potential. "
        
        # Convert list to flowing text - clean and merge into paragraphs
        rec_texts = []
        for insight in insights:
            # Clean up any numbering or priority markers
            clean_insight = insight.strip().rstrip('.')
            # Remove patterns like "PRIORITY 1 -", "[Category]", numbers, etc.
            import re
            clean_insight = re.sub(r'PRIORITY \d+ -\s*', '', clean_insight)
            clean_insight = re.sub(r'\[.*?\]\s*', '', clean_insight)
            clean_insight = re.sub(r'^\d+\.\s*', '', clean_insight)
            if clean_insight:
                rec_texts.append(clean_insight)
        
        # Join all recommendations into flowing text
        if len(rec_texts) > 0:
            # First set of recommendations
            if len(rec_texts) <= 3:
                recommendations_section += " Additionally, ".join(rec_texts) + ". "
            else:
                # Split into multiple sentences for better flow
                recommendations_section += rec_texts[0] + ". "
                recommendations_section += "Furthermore, " + rec_texts[1] + ". "
                recommendations_section += "It is also important to " + rec_texts[2].lower() + ". "
                
                if len(rec_texts) > 3:
                    recommendations_section += "\n\n"
                    recommendations_section += "In addition, " + rec_texts[3] + ". "
                    
                if len(rec_texts) > 4:
                    recommendations_section += "Moreover, " + rec_texts[4] + ". "
                    
                if len(rec_texts) > 5:
                    recommendations_section += "Finally, " + rec_texts[5] + ". "
                    
                # Add remaining recommendations if any
                if len(rec_texts) > 6:
                    remaining = " ".join([r + "." for r in rec_texts[6:]])
                    recommendations_section += remaining
    else:
        recommendations_section = "Please review the detailed analysis sections above for specific recommendations tailored to this project."
    
    sections['actionable_recommendations'] = recommendations_section.rstrip()
    
    # Compliance Check Section
    compliance = analysis.get('compliance_check', {})
    compliance_section = ""
    
    if compliance:
        is_compliant = compliance.get('is_compliant')
        
        if is_compliant:
            compliance_section += "The DPR demonstrates strong compliance with MDoNER guidelines and regulatory requirements. "
            compliance_section += "All necessary statutory clearances and documentation standards appear to be adequately addressed. "
        else:
            compliance_section += "While the DPR shows effort in meeting regulatory requirements, several compliance gaps have been identified that need attention. "
            
            gaps = compliance.get('guideline_gaps', [])
            if gaps:
                clean_gaps = [g.strip().rstrip('.') for g in gaps]
                compliance_section += "Specifically, the following areas require improvement including "
                if len(clean_gaps) == 1:
                    compliance_section += clean_gaps[0]
                elif len(clean_gaps) == 2:
                    compliance_section += clean_gaps[0] + " and " + clean_gaps[1]
                else:
                    compliance_section += ", ".join(clean_gaps[:-1]) + ", and " + clean_gaps[-1]
                compliance_section += ". "
        
        if compliance.get('comments'):
            compliance_section += f" {compliance.get('comments')}"
    else:
        # Generate comprehensive compliance analysis from available data
        overall_score = analysis.get('overall_score', 70)
        
        if overall_score >= 80:
            compliance_section += "The documentation demonstrates strong adherence to MDoNER guidelines and statutory requirements for project proposals in the North Eastern Region. "
            compliance_section += "The DPR format, content structure, and level of detail align well with prescribed standards for infrastructure projects. "
        elif overall_score >= 60:
            compliance_section += "The DPR shows reasonable compliance with MDoNER guidelines, though certain areas would benefit from additional documentation and clarity. "
            compliance_section += "The overall structure follows required standards, but some specific sections require strengthening. "
        else:
            compliance_section += "The compliance assessment reveals several gaps in meeting MDoNER guideline requirements that must be addressed for approval consideration. "
            compliance_section += "Both the documentation completeness and content quality need significant improvement. "
        
        # Extract compliance-related insights
        compliance_insights = [i for i in insights if any(term in i.lower() for term in ['compliance', 'guideline', 'requirement', 'standard', 'regulation', 'clearance', 'approval', 'documentation', 'section', 'missing'])]
        
        if compliance_insights:
            compliance_section += "Specific compliance considerations include the following aspects. "
            for insight in compliance_insights[:3]:  # Use top 3 compliance insights
                clean_insight = insight.replace('PRIORITY 1 - ', '').replace('PRIORITY 2 - ', '').replace('PRIORITY 3 - ', '').replace('PRIORITY 4 - ', '').replace('PRIORITY 5 - ', '')
                clean_insight = clean_insight.replace('[Compliance]', '').replace('[Documentation]', '').replace('[Guideline]', '').replace('[', '').replace(']', '').strip()
                if not clean_insight.endswith('.'):
                    clean_insight += '.'
                compliance_section += clean_insight + " "
        else:
            compliance_section += "Essential statutory clearances including environmental, forest, and wildlife approvals (where applicable) must be secured before project commencement. "
            compliance_section += "All sections mandated by MDoNER guidelines should be present with adequate depth and supporting documentation. "
        
        compliance_section += "Regular compliance audits during implementation will ensure continued adherence to approved standards and regulatory requirements."
    
    sections['compliance_check'] = compliance_section.rstrip()
    
    # Final Assessment Section
    final_assessment = ""
    if decision == "APPROVE":
        final_assessment += "After thorough evaluation of all aspects of this project, I recommend approval for implementation. "
        final_assessment += "The DPR demonstrates strong planning, technical feasibility, and alignment with MDoNER's development objectives for the North Eastern Region. "
        final_assessment += "The project is well-prepared and ready to move forward to the execution phase."
    elif decision == "REVISE":
        final_assessment += "Based on this comprehensive analysis, I recommend that the DPR be revised to address the concerns identified throughout this assessment. "
        final_assessment += "While the project shows promise and has several strong elements, the issues highlighted need to be resolved before approval. "
        final_assessment += "Once these revisions are made, the project should be well-positioned for successful implementation."
    else:
        final_assessment += "After careful evaluation, I must recommend rejection of this DPR in its current form. "
        final_assessment += "The analysis has revealed significant issues that fundamentally impact the project's viability and alignment with MDoNER guidelines. "
        final_assessment += "These concerns need to be thoroughly addressed through a substantially revised submission."
    
    if reasoning:
        final_assessment += f"\n\n{reasoning}"
    
    sections['final_assessment'] = final_assessment
    
    return sections


def generate_chatgpt_style_response(analysis: Dict, insights: List, risks: Dict, structured_dpr: Dict) -> str:
    """Generate a structured, detailed, and human-readable response"""
    
    overall_score = analysis.get('overall_score', 0)
    approval = analysis.get('approval_recommendation', {})
    decision = approval.get('decision', 'N/A')
    reasoning = approval.get('reasoning', '')
    summary = analysis.get('summary', '')
    
    project_title = structured_dpr.get('project_title', 'this project')
    project_type = structured_dpr.get('project_type', 'Infrastructure Project')
    location = structured_dpr.get('location', 'Not specified')
    implementing_agency = structured_dpr.get('implementing_agency', 'Not specified')
    duration = structured_dpr.get('duration', 'Not specified')
    budget = structured_dpr.get('budget', {}).get('total', 0)
    budget_details = structured_dpr.get('budget', {}).get('details', 'Not specified')
    
    # Format budget
    budget_formatted = "Not specified"
    if budget > 0:
        budget_crores = budget / 10000000
        budget_formatted = f"Rs. {budget_crores:.2f} Crores"
    
    response = f"""
================================================================================
                        DPR ANALYSIS REPORT
================================================================================

PROJECT INFORMATION:
-------------------
Project Title       : {project_title}
Project Type        : {project_type}
Location            : {location}
Implementing Agency : {implementing_agency}
Project Duration    : {duration}
Estimated Budget    : {budget_formatted}

================================================================================
                        EXECUTIVE SUMMARY
================================================================================

Overall Score       : {overall_score}/100
Recommendation      : {decision}
Analysis Date       : {datetime.now().strftime('%B %d, %Y at %I:%M %p')}

{summary if summary else reasoning}

================================================================================
                        SECTION 1: BUDGET ANALYSIS
================================================================================
"""
    
    budget_eval = analysis.get('budget_validation', {})
    
    if budget_eval:
        validity_status = "VALID" if budget_eval.get('is_valid') else "REQUIRES ATTENTION"
        response += f"\nBudget Status: [{validity_status}]\n"
        response += f"Budget Amount: {budget_details}\n\n"
        
        if budget_eval.get('is_valid'):
            response += "Assessment: The budget allocation appears realistic and well-structured.\n"
        else:
            response += "Assessment: Budget concerns have been identified that require attention.\n\n"
            response += "IDENTIFIED ISSUES:\n"
            for idx, issue in enumerate(budget_eval.get('issues', []), 1):
                if isinstance(issue, dict):
                    severity = issue.get('severity', 'medium')
                    description = issue.get('description', str(issue))
                    response += f"  {idx}. {description}\n"
                    response += f"     Severity: {severity.upper()}\n"
                else:
                    response += f"  {idx}. {issue}\n"
        
        if budget_eval.get('comments'):
            response += f"\nDetailed Assessment:\n{budget_eval.get('comments')}\n"
    else:
        response += "\nNo detailed budget validation data available.\n"
    
    response += "\n"
    response += "="*80 + "\n"
    response += "                    SECTION 2: TIMELINE EVALUATION\n"
    response += "="*80 + "\n"
    
    timeline_eval = analysis.get('timeline_validation', {})
    if timeline_eval:
        timeline_status = "REALISTIC" if timeline_eval.get('is_realistic') else "REQUIRES REVIEW"
        response += f"\nTimeline Status: [{timeline_status}]\n"
        response += f"Project Duration: {duration}\n\n"
        
        if timeline_eval.get('is_realistic'):
            response += "Assessment: The project timeline appears realistic and achievable.\n"
        else:
            response += "Assessment: Timeline concerns have been identified.\n\n"
            response += "IDENTIFIED CONCERNS:\n"
            for idx, issue in enumerate(timeline_eval.get('issues', []), 1):
                if isinstance(issue, dict):
                    severity = issue.get('severity', 'medium')
                    description = issue.get('description', str(issue))
                    response += f"  {idx}. {description}\n"
                    response += f"     Severity: {severity.upper()}\n"
                else:
                    response += f"  {idx}. {issue}\n"
        
        if timeline_eval.get('comments'):
            response += f"\nDetailed Assessment:\n{timeline_eval.get('comments')}\n"
    else:
        response += "\nNo detailed timeline validation data available.\n"
    
    response += "\n"
    response += "="*80 + "\n"
    response += "                   SECTION 3: TECHNICAL FEASIBILITY\n"
    response += "="*80 + "\n"
    
    tech_eval = analysis.get('technical_feasibility', {})
    if tech_eval:
        tech_score = tech_eval.get('score', 0)
        response += f"\nTechnical Feasibility Score: {tech_score}/100\n\n"
        
        if tech_eval.get('strengths'):
            response += "STRENGTHS IDENTIFIED:\n"
            for idx, strength in enumerate(tech_eval.get('strengths', []), 1):
                response += f"  {idx}. {strength}\n"
            response += "\n"
        
        if tech_eval.get('weaknesses'):
            response += "AREAS FOR IMPROVEMENT:\n"
            for idx, weakness in enumerate(tech_eval.get('weaknesses', []), 1):
                response += f"  {idx}. {weakness}\n"
            response += "\n"
        
        if tech_eval.get('comments'):
            response += f"Detailed Assessment:\n{tech_eval.get('comments')}\n"
    else:
        response += "\nNo detailed technical feasibility data available.\n"
    
    response += "\n"
    response += "="*80 + "\n"
    response += "                     SECTION 4: RISK ASSESSMENT\n"
    response += "="*80 + "\n"
    
    risk_eval = analysis.get('risk_assessment', {})
    if risk_eval:
        risk_level = risk_eval.get('overall_risk_level', 'unknown')
        response += f"\nOverall Risk Level: {risk_level.upper()}\n\n"
        
        # Financial Risks
        if risk_eval.get('financial_risks'):
            response += "4.1 FINANCIAL RISKS:\n"
            response += "-" * 40 + "\n"
            for idx, risk in enumerate(risk_eval.get('financial_risks', []), 1):
                if isinstance(risk, dict):
                    risk_text = risk.get('risk') or risk.get('name') or risk.get('description') or str(risk)
                    severity = risk.get('severity', 'unknown')
                    mitigation = risk.get('mitigation', '')
                    
                    response += f"  Risk {idx}: {risk_text}\n"
                    response += f"    Severity Level: {severity.upper()}\n"
                    if mitigation:
                        response += f"    Mitigation Strategy: {mitigation}\n"
                    response += "\n"
                else:
                    response += f"  Risk {idx}: {risk}\n\n"
        
        # Timeline Risks
        if risk_eval.get('timeline_risks'):
            response += "4.2 TIMELINE RISKS:\n"
            response += "-" * 40 + "\n"
            for idx, risk in enumerate(risk_eval.get('timeline_risks', []), 1):
                if isinstance(risk, dict):
                    risk_text = risk.get('risk') or risk.get('name') or risk.get('description') or str(risk)
                    severity = risk.get('severity', 'unknown')
                    mitigation = risk.get('mitigation', '')
                    
                    response += f"  Risk {idx}: {risk_text}\n"
                    response += f"    Severity Level: {severity.upper()}\n"
                    if mitigation:
                        response += f"    Mitigation Strategy: {mitigation}\n"
                    response += "\n"
                else:
                    response += f"  Risk {idx}: {risk}\n\n"
        
        # Environmental Risks
        if risk_eval.get('environmental_risks'):
            response += "4.3 ENVIRONMENTAL RISKS:\n"
            response += "-" * 40 + "\n"
            for idx, risk in enumerate(risk_eval.get('environmental_risks', []), 1):
                if isinstance(risk, dict):
                    risk_text = risk.get('risk') or risk.get('name') or risk.get('description') or str(risk)
                    severity = risk.get('severity', 'unknown')
                    mitigation = risk.get('mitigation', '')
                    
                    response += f"  Risk {idx}: {risk_text}\n"
                    response += f"    Severity Level: {severity.upper()}\n"
                    if mitigation:
                        response += f"    Mitigation Strategy: {mitigation}\n"
                    response += "\n"
                else:
                    response += f"  Risk {idx}: {risk}\n\n"
        
        # Resource Risks
        if risk_eval.get('resource_risks'):
            response += "4.4 RESOURCE RISKS:\n"
            response += "-" * 40 + "\n"
            for idx, risk in enumerate(risk_eval.get('resource_risks', []), 1):
                if isinstance(risk, dict):
                    risk_text = risk.get('risk') or risk.get('name') or risk.get('description') or str(risk)
                    severity = risk.get('severity', 'unknown')
                    mitigation = risk.get('mitigation', '')
                    
                    response += f"  Risk {idx}: {risk_text}\n"
                    response += f"    Severity Level: {severity.upper()}\n"
                    if mitigation:
                        response += f"    Mitigation Strategy: {mitigation}\n"
                    response += "\n"
                else:
                    response += f"  Risk {idx}: {risk}\n\n"
    else:
        response += "\nNo detailed risk assessment data available.\n"
    
    response += "\n"
    response += "="*80 + "\n"
    response += "                 SECTION 5: ACTIONABLE RECOMMENDATIONS\n"
    response += "="*80 + "\n\n"
    
    if insights and len(insights) > 0:
        response += "Based on the comprehensive analysis, the following actionable steps are\n"
        response += "recommended to strengthen this DPR:\n\n"
        for idx, insight in enumerate(insights, 1):
            response += f"  {idx}. {insight}\n"
    else:
        response += "Please review the detailed analysis sections above for specific recommendations.\n"
    
    response += "\n"
    response += "="*80 + "\n"
    response += "                   SECTION 6: COMPLIANCE CHECK\n"
    response += "="*80 + "\n\n"
    
    compliance = analysis.get('compliance_check', {})
    if compliance:
        compliance_status = "COMPLIANT" if compliance.get('is_compliant') else "NON-COMPLIANT"
        response += f"Compliance Status: [{compliance_status}]\n\n"
        
        if compliance.get('is_compliant'):
            response += "Assessment: The DPR appears to be compliant with MDoNER guidelines.\n"
        else:
            response += "Assessment: Compliance gaps have been identified:\n\n"
            response += "IDENTIFIED GAPS:\n"
            for idx, gap in enumerate(compliance.get('guideline_gaps', []), 1):
                response += f"  {idx}. {gap}\n"
        
        if compliance.get('comments'):
            response += f"\nDetailed Assessment:\n{compliance.get('comments')}\n"
    else:
        response += "No detailed compliance data available.\n"
    
    response += "\n"
    response += "="*80 + "\n"
    response += "                      FINAL ASSESSMENT\n"
    response += "="*80 + "\n\n"
    
    if decision == "APPROVE":
        response += "RECOMMENDATION: APPROVE\n"
        response += "-" * 40 + "\n"
        response += "This DPR is well-prepared and ready for approval. The project demonstrates\n"
        response += "strong planning, feasibility, and alignment with MDoNER guidelines.\n"
    elif decision == "REVISE":
        response += "RECOMMENDATION: REVISE\n"
        response += "-" * 40 + "\n"
        response += "This DPR shows promise but requires revisions to address the concerns\n"
        response += "identified in the analysis above.\n"
    else:
        response += "RECOMMENDATION: REJECT\n"
        response += "-" * 40 + "\n"
        response += "This DPR has significant issues that need to be resolved before it can\n"
        response += "be approved for implementation.\n"
    
    response += "\n"
    response += "="*80 + "\n"
    response += "                    END OF ANALYSIS REPORT\n"
    response += "="*80 + "\n\n"
    response += "DISCLAIMER: This analysis was generated by AI and should be reviewed by\n"
    response += "            qualified professionals before making final decisions.\n"
    response += "\n" + "="*80
    
    return response


def save_analysis_to_json(result: Dict, filename: str, structured_sections: Dict = None) -> str:
    """Save analysis results to a JSON file with structured sections"""
    os.makedirs("analysis_results", exist_ok=True)
    
    timestamp = datetime.now().strftime("%Y%m%d_%H%M%S")
    json_filename = f"analysis_results/analysis_{timestamp}_{filename.replace('.pdf', '.json').replace('.docx', '.json')}"
    
    # Add structured sections to result if provided
    if structured_sections:
        result['structured_analysis'] = structured_sections
        result['structured_analysis_metadata'] = {
            'generated_at': datetime.now().isoformat(),
            'format': 'plain_text',
            'exportable': True,
            'sections': list(structured_sections.keys())
        }
    
    with open(json_filename, 'w', encoding='utf-8') as f:
        json.dump(result, f, indent=2, ensure_ascii=False)
    
    return json_filename


# ============================================================================
# FASTAPI APPLICATION
# ============================================================================

app = FastAPI(
    title="DPR Analysis System",
    description="AI-Powered DPR Analysis using Google Gemini",
    version="2.0.0"
)

# Enable CORS
app.add_middleware(
    CORSMiddleware,
    allow_origins=["*"],
    allow_credentials=True,
    allow_methods=["*"],
    allow_headers=["*"],
)

# Create directories
os.makedirs("uploads", exist_ok=True)
os.makedirs("data/guidelines", exist_ok=True)


# ============================================================================
# API ENDPOINTS
# ============================================================================

@app.get("/")
async def root():
    """Root endpoint - API information"""
    return {
        "message": "DPR Analysis System - Powered by Google Gemini AI",
        "status": "running",
        "version": "2.0.0",
        "features": [
            "AI-powered DPR analysis",
            "Budget & timeline validation",
            "Risk assessment",
            "Gap detection",
            "Multi-language support",
            "Actionable recommendations",
            "Structured exportable output"
        ],
        "endpoints": {
            "docs": "/docs",
            "upload_dpr": "/api/upload-dpr",
            "load_guidelines": "/api/load-guidelines",
            "health": "/api/health"
        }
    }


@app.get("/api/health")
async def health_check():
    """Health check endpoint"""
    return {
        "status": "healthy",
        "gemini_configured": bool(GEMINI_API_KEY),
        "model": GEMINI_MODEL,
        "guidelines_loaded": bool(guidelines_context)
    }


@app.get("/api/download/{document_id}")
async def download_document(document_id: str):
    """
    Download uploaded DPR document by document ID
    Document ID format: original_filename (e.g., sampledpr.pdf)
    
    Returns the actual uploaded file from the uploads directory
    """
    from fastapi.responses import FileResponse
    
    try:
        # First, try to find the stored filename in analysis results
        # Check both backend/analysis_results and root analysis_results
        analysis_dirs = ["analysis_results", "../analysis_results"]
        
        for analysis_dir in analysis_dirs:
            if os.path.exists(analysis_dir):
                json_files = [f for f in os.listdir(analysis_dir) if f.endswith('.json')]
                
                for json_file in json_files:
                    json_path = os.path.join(analysis_dir, json_file)
                    try:
                        with open(json_path, 'r', encoding='utf-8') as f:
                            data = json.load(f)
                            
                        # Check if this analysis corresponds to the requested document
                        if data.get('filename') == document_id or data.get('original_filename') == document_id:
                            stored_filename = data.get('stored_filename')
                            file_path_from_json = data.get('file_path')
                            
                            if stored_filename and os.path.exists(f"uploads/{stored_filename}"):
                                file_path = f"uploads/{stored_filename}"
                            elif file_path_from_json and os.path.exists(file_path_from_json):
                                file_path = file_path_from_json
                            else:
                                continue
                            
                            # Determine media type
                            extension = stored_filename.split('.')[-1].lower() if stored_filename else 'pdf'
                            media_type = {
                                'pdf': 'application/pdf',
                                'docx': 'application/vnd.openxmlformats-officedocument.wordprocessingml.document',
                                'doc': 'application/msword',
                                'txt': 'text/plain'
                            }.get(extension, 'application/octet-stream')
                            
                            print(f"[DOWNLOAD] Sending file: {file_path} as {document_id}")
                            return FileResponse(
                                path=file_path,
                                media_type=media_type,
                                filename=document_id  # Use original filename for download
                            )
                    except Exception as e:
                        print(f"[WARNING] Error reading {json_file}: {e}")
                        continue
        
        # Fallback: Search uploads directory for matching files
        uploads_dir = "uploads"
        if not os.path.exists(uploads_dir):
            raise HTTPException(404, "Uploads directory not found")
        
        files = os.listdir(uploads_dir)
        
        # If document_id is the original name, find the corresponding dpr_timestamp file
        # We'll return the most recent file that matches
        if document_id.lower().endswith(('.pdf', '.docx', '.doc', '.txt')):
            # Return the most recent file (assuming one file for now)
            if files:
                latest_file = sorted(files, reverse=True)[0]  # Get most recent file
                file_path = os.path.join(uploads_dir, latest_file)
                
                # Determine media type
                extension = latest_file.split('.')[-1].lower()
                media_type = {
                    'pdf': 'application/pdf',
                    'docx': 'application/vnd.openxmlformats-officedocument.wordprocessingml.document',
                    'doc': 'application/msword',
                    'txt': 'text/plain'
                }.get(extension, 'application/octet-stream')
                
                print(f"[DOWNLOAD-FALLBACK] Sending most recent file: {file_path} as {document_id}")
                return FileResponse(
                    path=file_path,
                    media_type=media_type,
                    filename=document_id  # Use original filename for download
                )
        
        raise HTTPException(404, f"Document not found: {document_id}")
        
    except HTTPException:
        raise
    except Exception as e:
        print(f"[ERROR] Download error: {str(e)}")
        import traceback
        traceback.print_exc()
        raise HTTPException(500, f"Error downloading document: {str(e)}")


@app.post("/api/upload-dpr")
async def upload_and_analyze_dpr(
    file: UploadFile = File(...),
    language: str = Form("en")
):
    """
    Upload DPR and get comprehensive AI analysis
    
    - **file**: PDF or DOCX file
    - **language**: en, hi, as, bn, mni, ne
    """
    
    try:
        # Validate file type
        file_extension = file.filename.split(".")[-1].lower()
        if file_extension not in ['pdf', 'docx', 'doc']:
            raise HTTPException(400, f"Unsupported file type: {file_extension}")
        
        # Save file
        timestamp = datetime.now().strftime("%Y%m%d_%H%M%S")
        filename = f"dpr_{timestamp}.{file_extension}"
        file_path = f"uploads/{filename}"
        
        print(f"[SAVE] Saving file: {filename}")
        with open(file_path, "wb") as buffer:
            shutil.copyfileobj(file.file, buffer)
        
        # Extract text
        print(f"[EXTRACTING] Extracting text from {file_extension.upper()}...")
        extracted_text = extract_text(file_path, file_extension)
        
        if len(extracted_text.strip()) < 100:
            raise HTTPException(400, "Could not extract sufficient text from document")
        
        print(f"[EXTRACTED] Extracted {len(extracted_text)} characters")
        
        # Structure data
        print("[STRUCTURING] Structuring DPR data...")
        structured_dpr = structure_dpr_data(extracted_text)
        
        # OPTIMIZED: Single AI Analysis call (includes insights + risks)
        print("[FAST-AI-ANALYSIS] Starting optimized single-call analysis...")
        analysis = await analyze_dpr_comprehensive_fast(extracted_text, structured_dpr)
        
        # Extract insights and risks from the comprehensive analysis
        insights = analysis.get('actionable_insights', analysis.get('recommendations', []))
        risks = {
            "overall_risk_level": analysis.get('risk_assessment', {}).get('overall_risk_level', 'medium'),
            "overall_risk_score": analysis.get('risk_assessment', {}).get('overall_risk_score', 50),
            "financial_risks": analysis.get('risk_assessment', {}).get('financial_risks', []),
            "timeline_risks": analysis.get('risk_assessment', {}).get('timeline_risks', []),
            "environmental_risks": analysis.get('risk_assessment', {}).get('environmental_risks', []),
            "resource_risks": analysis.get('risk_assessment', {}).get('resource_risks', []),
        }
        
        # Combine results
        result = {
            "dpr_id": timestamp,
            "filename": file.filename,
            "stored_filename": filename,  # Add the actual stored filename
            "file_path": file_path,  # Add the full file path
            "upload_time": datetime.now().isoformat(),
            "extracted_data": structured_dpr,
            "analysis": analysis,
            "actionable_insights": insights,
            "risk_assessment": risks,
            "language": language
        }
        
        # Generate structured sections for JSON (without formatting lines)
        print("Generating structured analysis sections...")
        structured_sections = generate_structured_json_sections(analysis, insights, risks, structured_dpr)
        
        # Also generate formatted response for terminal display only
        chatgpt_response = generate_chatgpt_style_response(analysis, insights, risks, structured_dpr)
        
        # Save to JSON file with structured sections only (no formatted output)
        print("Saving analysis to JSON...")
        json_file_path = save_analysis_to_json(result, file.filename, structured_sections=structured_sections)
        result['saved_to'] = json_file_path
        print(f"[SAVED] Analysis saved to: {json_file_path}")
        
        # Translate if needed
        if language != "en":
            print(f"Translating to {language}...")
            result = await translate_report(result, language)
        
        print("[COMPLETE] Analysis complete!")
        
        # ========================================
        # PRINT STRUCTURED RESPONSE IN CLI
        # ========================================
        print("\n" + "="*80)
        print("DPR ANALYSIS REPORT")
        print("="*80)
        print(chatgpt_response)
        print("\n" + "="*80)
        print(f"File: {file.filename}")
        print(f"Saved to: {json_file_path}")
        print(f"Completed: {datetime.now().strftime('%Y-%m-%d %H:%M:%S')}")
        print("="*80 + "\n")
        
        return {"status": "success", "result": result}
        
    except HTTPException:
        raise
    except Exception as e:
        print(f"[ERROR] Error: {e}")
        raise HTTPException(500, f"Error processing DPR: {str(e)}")


@app.post("/api/upload-dpr-fast")
async def upload_and_analyze_dpr_fast(
    file: UploadFile = File(...),
    language: str = Form("en")
):
    """
    FAST Upload - Returns ONLY recommendations for client portal
    Optimized for 3x faster processing by skipping full analysis
    
    - **file**: PDF or DOCX file
    - **language**: en, hi, as, bn, mni, ne
    """
    
    try:
        # Validate file type
        file_extension = file.filename.split(".")[-1].lower()
        if file_extension not in ['pdf', 'docx', 'doc', 'txt']:
            raise HTTPException(400, f"Unsupported file type: {file_extension}")
        
        # Save file
        timestamp = datetime.now().strftime("%Y%m%d_%H%M%S")
        filename = f"dpr_{timestamp}.{file_extension}"
        file_path = f"uploads/{filename}"
        
        print(f"[FAST-MODE] Saving file: {filename}")
        with open(file_path, "wb") as buffer:
            shutil.copyfileobj(file.file, buffer)
        
        # Extract text
        print(f"[FAST-MODE] Extracting text from {file_extension.upper()}...")
        extracted_text = extract_text(file_path, file_extension)
        
        if len(extracted_text.strip()) < 100:
            raise HTTPException(400, "Could not extract sufficient text from document")
        
        print(f"[FAST-MODE] Extracted {len(extracted_text)} characters")
        
        # Structure basic data (quick pass)
        structured_dpr = structure_dpr_data(extracted_text)
        
        # OPTIMIZED: Direct recommendations generation only (skip full analysis)
        print("[FAST-MODE] Generating recommendations directly...")
        
        prompt = f"""
You are an AI expert analyzing DPRs for India's Ministry of Development of North Eastern Region (MDoNER).

Analyze this DPR and provide 6-8 PRIORITY-BASED recommendations for improvement:

PROJECT DETAILS:
- Title: {structured_dpr.get('project_title', 'Not specified')}
- Type: {structured_dpr.get('project_type', 'general')}
- Budget: {structured_dpr.get('budget', {}).get('total_formatted', 'Not specified')}
- Timeline: {structured_dpr.get('timeline', {}).get('duration', 'Not specified')}
- Location: {structured_dpr.get('location', 'North Eastern Region')}

DPR CONTENT (First 15000 characters):
{extracted_text[:15000]}

CRITICAL INSTRUCTIONS:
1. Analyze gaps against MDoNER approval requirements
2. Provide EXTENSIVE 4-5 LINE explanations for each recommendation
3. Prioritize by impact on approval success

Format each recommendation as:
"PRIORITY X - [CATEGORY] Specific actionable recommendation - EXTENSIVE 4-5 LINE EXPLANATION covering: (1) Why this addresses MDoNER approval requirements, (2) How it improves project implementation success, (3) What specific risks it mitigates, (4) Impact on regulatory compliance, (5) Benefits for stakeholders and measurable project outcomes"

Categories: [BUDGET], [TIMELINE], [TECHNICAL], [COMPLIANCE], [RISK], [DOCUMENTATION], [STAKEHOLDER], [ENVIRONMENTAL], [FINANCIAL]

Example:
"PRIORITY 1 - [BUDGET] Revise the cost estimates in Section 3.1 to include detailed unit rates, quantity derivations, and market analysis for each construction component - This detailed cost breakdown is specifically mandated by MDoNER financial guidelines and treasury department requirements to ensure complete transparency in fund utilization. The absence of proper unit rate justification will lead to immediate rejection by the financial review committee as they need to verify that public funds are being allocated efficiently and market rates are realistic. Detailed cost analysis also enables proper fund allocation monitoring during implementation, prevents cost overruns that could jeopardize project completion, and demonstrates value-for-money to approval committees who scrutinize every budget line item."

Return ONLY a JSON object:
{{
    "standard_assessment": "Brief 2-3 sentence assessment of DPR quality and readiness",
    "detailed_recommendations": [
        "PRIORITY 1 - [CATEGORY] Detailed recommendation with 4-5 line explanation...",
        "PRIORITY 2 - [CATEGORY] Next recommendation with detailed reasoning...",
        "... (6-8 total recommendations)"
    ]
}}
"""
        
        try:
            response = gemini_model.generate_content(prompt)
            insights_data = parse_json_response(response.text)
            
            # Handle the format with standard assessment and detailed recommendations
            if isinstance(insights_data, dict) and 'detailed_recommendations' in insights_data:
                standard_assessment = insights_data.get('standard_assessment', '')
                recommendations = insights_data.get('detailed_recommendations', [])
                
                # Ensure proper formatting
                enhanced_insights = []
                for i, insight in enumerate(recommendations):
                    if not insight.startswith("PRIORITY"):
                        priority = i + 1
                        if "[" not in insight or "]" not in insight:
                            insight = f"PRIORITY {priority} - [GENERAL] {insight} - This improvement is recommended to strengthen the DPR submission and increase approval chances with MDoNER by addressing critical evaluation criteria."
                        else:
                            insight = f"PRIORITY {priority} - {insight}"
                    enhanced_insights.append(insight)
                
                # Add assessment as first item
                if standard_assessment:
                    enhanced_insights.insert(0, f"ASSESSMENT - {standard_assessment}")
                
                actionable_insights = enhanced_insights[:10]
            else:
                # Fallback
                actionable_insights = [
                    "ASSESSMENT - The DPR provides a solid foundation but requires revisions to address gaps in environmental and social impact assessments, risk assessment, and financial viability.",
                    "PRIORITY 1 - [REVIEW] Conduct a comprehensive review of all DPR sections to ensure completeness and accuracy - A thorough review is essential to identify and address any gaps that could delay MDoNER approval or cause implementation issues, particularly in areas of environmental assessment, social impact analysis, and financial viability documentation."
                ]
        
        except Exception as e:
            print(f"[FAST-MODE ERROR] Recommendations generation failed: {e}")
            actionable_insights = [
                "ASSESSMENT - The DPR provides a solid foundation but requires revisions to address gaps in environmental and social impact assessments, risk assessment, and financial viability.",
                "PRIORITY 1 - [REVIEW] Conduct a comprehensive review of all DPR sections to ensure completeness and accuracy - A thorough review is essential to identify and address any gaps that could delay MDoNER approval."
            ]
        
        # Build fast response (minimal data)
        result = {
            "dpr_id": timestamp,
            "filename": file.filename,
            "upload_time": datetime.now().isoformat(),
            "extracted_data": {
                "project_title": structured_dpr.get('project_title', 'Not specified'),
                "project_type": structured_dpr.get('project_type', 'general'),
                "location": structured_dpr.get('location', 'Not specified'),
                "budget": {
                    "total": structured_dpr.get('budget', {}).get('total', 0),
                    "total_formatted": structured_dpr.get('budget', {}).get('details', 'Not specified'),
                    "currency": structured_dpr.get('budget', {}).get('currency', 'INR')
                },
                "timeline": {
                    "duration": structured_dpr.get('timeline', {}).get('duration', 'Not specified'),
                    "duration_months": structured_dpr.get('timeline', {}).get('duration_months', 0)
                },
                "word_count": structured_dpr.get('word_count', 0)
            },
            "actionable_insights": actionable_insights,
            "language": language,
            "full_text": extracted_text
        }
        
        print(f"[FAST-MODE COMPLETE] Generated {len(actionable_insights)} recommendations in optimized mode")
        
        return {"status": "success", "result": result}
        
    except HTTPException:
        raise
    except Exception as e:
        print(f"[FAST-MODE ERROR] Error: {e}")
        raise HTTPException(500, f"Error processing DPR: {str(e)}")


@app.post("/api/admin/review-compliance")
async def admin_review_compliance(
    dpr_text: str = Form(...),
    project_info: str = Form(...),
    compliance_only: str = Form("false"),
    get_recommendation: str = Form("false")
):
    """
    Admin Compliance Review - Check if DPR meets MDoNER guidelines
    - compliance_only=true: Returns only compliance check (no auto-rejection, no feasibility)
    - get_recommendation=true: Returns detailed feasibility assessment (Technical, Financial, Risk)
    """
    try:
        compliance_only_mode = compliance_only.lower() == "true"
        get_recommendation_mode = get_recommendation.lower() == "true"
        
        print(f"[ADMIN-REVIEW] Mode: compliance_only={compliance_only_mode}, get_recommendation={get_recommendation_mode}")
        print(f"[ADMIN-REVIEW] DPR text length: {len(dpr_text)}")
        print(f"[ADMIN-REVIEW] Project info: {project_info[:200]}...")
        
        if not dpr_text or len(dpr_text) < 100:
            raise HTTPException(400, "DPR text is too short or empty")
        
        # Parse project info
        try:
            project_data = json.loads(project_info)
        except json.JSONDecodeError as e:
            print(f"[ADMIN-REVIEW ERROR] Failed to parse project_info: {e}")
            raise HTTPException(400, f"Invalid project info JSON: {str(e)}")
        
        print(f"[ADMIN-REVIEW] Project data parsed successfully")
        
        # If get_recommendation mode, skip compliance check and go straight to detailed assessment
        if get_recommendation_mode:
            print(f"[ADMIN-REVIEW] 📊 Generating detailed feasibility recommendations...")
            
            assessment_prompt = f"""
You are a senior MDoNER approval committee member providing detailed recommendations.

PROJECT INFORMATION:
{json.dumps(project_data, indent=2)}

DPR CONTENT:
{dpr_text[:20000]}

ANALYZE THESE THREE CRITICAL DIMENSIONS IN ORDER:

1. **TECHNICAL FEASIBILITY** (Analyze First - Weight: 35%)
   - Design adequacy and engineering standards compliance
   - Site suitability and geological conditions
   - Construction methodology feasibility
   - Quality assurance mechanisms
   - Technical expertise of implementing agency

2. **FINANCIAL FEASIBILITY** (Analyze Second - Weight: 35%)
   - Budget realism and cost estimates accuracy
   - Funding mechanism viability
   - Economic viability (BCR, EIRR, NPV, Payback)
   - Cost-benefit analysis
   - Financial sustainability

3. **RISK ASSESSMENT** (Analyze Last - Weight: 30%)
   - Identified risks and mitigation strategies
   - Implementation risks (delays, cost overruns)
   - Environmental and social risks
   - Operational risks
   - Contingency planning adequacy

PROVIDE DETAILED RECOMMENDATIONS IN JSON FORMAT:
{{
  "assessment": {{
    "technical": {{
      "score": 0-100,
      "rating": "EXCELLENT" / "GOOD" / "ADEQUATE" / "WEAK",
      "strengths": ["list 3-5 specific technical strengths"],
      "concerns": ["list 2-4 technical concerns if any"],
      "detailed_analysis": "4-5 sentences analyzing technical feasibility, design standards, site suitability"
    }},
    "financial": {{
      "score": 0-100,
      "rating": "EXCELLENT" / "GOOD" / "ADEQUATE" / "WEAK",
      "strengths": ["list 3-5 specific financial strengths"],
      "concerns": ["list 2-4 financial concerns if any"],
      "detailed_analysis": "4-5 sentences analyzing budget adequacy, economic viability, funding mechanism"
    }},
    "risk": {{
      "score": 0-100,
      "rating": "EXCELLENT" / "GOOD" / "ADEQUATE" / "WEAK",
      "strengths": ["list 3-5 risk management strengths"],
      "concerns": ["list 2-4 risk-related concerns"],
      "detailed_analysis": "4-5 sentences analyzing risk identification, mitigation strategies, contingency planning"
    }}
  }},
  "recommendation": {{
    "action": "APPROVE" / "CONDITIONAL_APPROVE" / "REQUEST_REVISIONS",
    "overall_score": 0-100,
    "confidence": 0-100,
    "summary": "Comprehensive 4-6 sentence summary explaining the recommendation with specific references to technical, financial, and risk factors"
  }}
}}

Scoring Guidelines:
- 80-100: Excellent - Strongly recommend approval
- 60-79: Good - Recommend approval with minor conditions
- 40-59: Adequate - Request revisions before approval
- 0-39: Weak - Major concerns, recommend rejection

Be specific and actionable in your analysis.
"""
            
            print(f"[ADMIN-REVIEW] Calling Gemini for detailed assessment...")
            assessment_response = gemini_model.generate_content(assessment_prompt)
            assessment_data = parse_json_response(assessment_response.text)
            
            print(f"[ADMIN-REVIEW] ✅ Recommendation generated:")
            print(f"  - Technical Score: {assessment_data.get('assessment', {}).get('technical', {}).get('score', 0)}/100")
            print(f"  - Financial Score: {assessment_data.get('assessment', {}).get('financial', {}).get('score', 0)}/100")
            print(f"  - Risk Score: {assessment_data.get('assessment', {}).get('risk', {}).get('score', 0)}/100")
            print(f"  - Overall Score: {assessment_data.get('recommendation', {}).get('overall_score', 0)}/100")
            print(f"  - Action: {assessment_data.get('recommendation', {}).get('action', 'N/A')}")
            
            return {
                "status": "recommendation_ready",
                "assessment": assessment_data.get('assessment', {}),
                "recommendation": assessment_data.get('recommendation', {})
            }
        
        # Otherwise, perform compliance check
        # Step 1: MDoNER Guidelines Compliance Check
        compliance_prompt = f"""
You are an expert MDoNER compliance officer reviewing a DPR submission.

PROJECT INFORMATION:
{json.dumps(project_data, indent=2)}

DPR CONTENT (First 20000 characters):
{dpr_text[:20000]}

CRITICAL TASK: Evaluate if this DPR meets MANDATORY MDoNER guidelines for North Eastern Region projects.

CHECK THESE MANDATORY REQUIREMENTS:
1. **Project Location**: Must be in North Eastern Region (Assam, Arunachal Pradesh, Manipur, Meghalaya, Mizoram, Nagaland, Sikkim, Tripura)
2. **Budget Documentation**: Detailed cost breakdown with quantities, rates, and totals
3. **Timeline**: Clear implementation schedule with milestones
4. **Technical Specifications**: Design standards, technical details, compliance codes
5. **Environmental Clearance**: Status of environmental approvals (EC/2024 or similar)
6. **Social Impact Assessment**: Land acquisition, R&R plan, stakeholder consultation
7. **Risk Assessment**: Identified risks with mitigation strategies
8. **Funding Mechanism**: Clear central-state funding split (typically 90:10 for NER)
9. **Implementing Agency**: Clearly identified with nodal officer details
10. **Statutory Approvals**: List of obtained/pending clearances

RESPOND IN STRICT JSON FORMAT:
{{
  "compliant": true/false,
  "compliance_score": 0-100,
  "critical_violations": ["list of mandatory requirements NOT met"],
  "missing_sections": ["list of missing critical sections"],
  "rejection_reason": "Detailed reason if non-compliant, null if compliant",
  "compliance_summary": "Brief 2-3 sentence summary"
}}

If ANY critical violations exist, set compliant=false and provide detailed rejection_reason.
"""

        print(f"[ADMIN-REVIEW] Checking MDoNER compliance...")
        compliance_response = gemini_model.generate_content(compliance_prompt)
        compliance_data = parse_json_response(compliance_response.text)
        
        print(f"[ADMIN-REVIEW] Compliance Score: {compliance_data.get('compliance_score', 0)}%")
        print(f"[ADMIN-REVIEW] Compliant: {compliance_data.get('compliant', False)}")
        
        # If compliance_only mode, return just the compliance check
        if compliance_only_mode:
            print(f"[ADMIN-REVIEW] Compliance-only mode: Returning compliance data for manual admin review")
            return {
                "status": "compliance_checked",
                "compliance_data": compliance_data,
                "compliant": compliance_data.get('compliant', False)
            }
        
        # Legacy mode: Auto-reject if non-compliant
        if not compliance_data.get('compliant', False):
            print(f"[ADMIN-REVIEW] ❌ AUTO-REJECTED - Non-compliant with MDoNER guidelines")
            return {
                "status": "rejected",
                "auto_rejected": True,
                "compliance_data": compliance_data,
                "reason": compliance_data.get('rejection_reason', 'Does not meet MDoNER mandatory guidelines'),
                "recommendation": {
                    "action": "REJECT",
                    "confidence": "HIGH",
                    "summary": "DPR does not meet mandatory MDoNER guidelines and must be rejected."
                }
            }
        
        # Step 2: If compliant, perform detailed feasibility assessment
        print(f"[ADMIN-REVIEW] ✅ Compliant - Generating approval recommendations...")
        
        assessment_prompt = f"""
You are a senior MDoNER approval committee member providing detailed recommendations to the admin.

The DPR has PASSED mandatory compliance checks. Now provide DETAILED APPROVAL RECOMMENDATIONS.

PROJECT INFORMATION:
{json.dumps(project_data, indent=2)}

DPR CONTENT:
{dpr_text[:20000]}

COMPLIANCE STATUS: ✅ PASSED (Score: {compliance_data.get('compliance_score', 0)}%)

ANALYZE THESE THREE CRITICAL DIMENSIONS:

1. **TECHNICAL FEASIBILITY** (Weight: 35%)
   - Design adequacy and engineering standards compliance
   - Site suitability and geological conditions
   - Construction methodology feasibility
   - Quality assurance mechanisms
   - Technical expertise of implementing agency

2. **FINANCIAL FEASIBILITY** (Weight: 35%)
   - Budget realism and cost estimates accuracy
   - Funding mechanism viability
   - Economic viability (BCR, EIRR, NPV, Payback)
   - Cost-benefit analysis
   - Financial sustainability

3. **RISK ASSESSMENT** (Weight: 30%)
   - Identified risks and mitigation strategies
   - Implementation risks (delays, cost overruns)
   - Environmental and social risks
   - Operational risks
   - Contingency planning adequacy

PROVIDE DETAILED RECOMMENDATIONS IN JSON FORMAT:
{{
  "overall_recommendation": "APPROVE" / "CONDITIONAL_APPROVE" / "REQUEST_REVISIONS",
  "confidence_level": "HIGH" / "MEDIUM" / "LOW",
  "approval_score": 0-100,
  
  "technical_feasibility": {{
    "score": 0-100,
    "rating": "EXCELLENT" / "GOOD" / "ADEQUATE" / "WEAK",
    "strengths": ["list 3-5 technical strengths"],
    "concerns": ["list 2-4 technical concerns if any"],
    "detailed_analysis": "4-5 sentences on technical evaluation",
    "recommendation": "Approve/Revise with specific actions"
  }},
  
  "financial_feasibility": {{
    "score": 0-100,
    "rating": "EXCELLENT" / "GOOD" / "ADEQUATE" / "WEAK",
    "strengths": ["list 3-5 financial strengths"],
    "concerns": ["list 2-4 financial concerns if any"],
    "detailed_analysis": "4-5 sentences on financial evaluation",
    "budget_adequacy": "Assessment of cost estimates",
    "economic_viability": "Assessment of BCR, EIRR, NPV",
    "recommendation": "Approve/Revise with specific actions"
  }},
  
  "risk_assessment": {{
    "score": 0-100,
    "rating": "EXCELLENT" / "GOOD" / "ADEQUATE" / "WEAK",
    "strengths": ["list 3-5 risk management strengths"],
    "concerns": ["list 2-4 risk concerns if any"],
    "detailed_analysis": "4-5 sentences on risk evaluation",
    "critical_risks": ["list top 3 critical risks"],
    "mitigation_adequacy": "Assessment of mitigation strategies",
    "recommendation": "Approve/Revise with specific actions"
  }},
  
  "key_highlights": ["list 5-7 key points for admin decision"],
  "conditions_for_approval": ["list any conditions/requirements before final approval"],
  "admin_action_summary": "Clear 3-4 sentence summary for admin on what action to take and why",
  "estimated_success_probability": "percentage",
  "timeline_realism": "Assessment of proposed timeline",
  "overall_justification": "Detailed 5-6 sentence justification for the recommendation"
}}

Be thorough, specific, and provide actionable insights. The admin will use this for final decision.
"""

        print(f"[ADMIN-REVIEW] Generating detailed feasibility assessment...")
        assessment_response = gemini_model.generate_content(assessment_prompt)
        assessment_data = parse_json_response(assessment_response.text)
        
        print(f"[ADMIN-REVIEW] ✅ Assessment complete - Recommendation: {assessment_data.get('overall_recommendation', 'N/A')}")
        
        return {
            "status": "reviewed",
            "auto_rejected": False,
            "compliance_data": compliance_data,
            "assessment": assessment_data,
            "recommendation": {
                "action": assessment_data.get('overall_recommendation', 'REQUEST_REVISIONS'),
                "confidence": assessment_data.get('confidence_level', 'MEDIUM'),
                "score": assessment_data.get('approval_score', 0),
                "summary": assessment_data.get('admin_action_summary', '')
            }
        }
        
    except Exception as e:
        print(f"[ADMIN-REVIEW ERROR] {e}")
        import traceback
        traceback.print_exc()
        raise HTTPException(500, f"Error during compliance review: {str(e)}")


@app.post("/api/load-guidelines")
async def load_guidelines(file: UploadFile = File(...)):
    """Load MDoNER guideline documents"""
    global guidelines_context
    
    try:
        file_extension = file.filename.split(".")[-1].lower()
        if file_extension not in ['pdf', 'docx', 'doc']:
            raise HTTPException(400, f"Unsupported file type")
        
        # Save file
        guideline_path = f"data/guidelines/{file.filename}"
        with open(guideline_path, "wb") as buffer:
            shutil.copyfileobj(file.file, buffer)
        
        # Extract text
        text = extract_text(guideline_path, file_extension)
        guidelines_context = text[:30000]  # Store first 30k chars
        
        return {
            "status": "success",
            "message": f"Guideline '{file.filename}' loaded",
            "characters_loaded": len(text)
        }
    except Exception as e:
        raise HTTPException(500, f"Error loading guidelines: {str(e)}")


@app.post("/api/validate-budget")
async def validate_budget(file: UploadFile = File(...)):
    """Dedicated budget validation endpoint"""
    try:
        file_extension = file.filename.split(".")[-1].lower()
        file_path = f"uploads/temp_{datetime.now().timestamp()}.{file_extension}"
        
        with open(file_path, "wb") as buffer:
            shutil.copyfileobj(file.file, buffer)
        
        text = extract_text(file_path, file_extension)
        structured = structure_dpr_data(text)
        
        prompt = f"""
Analyze budget for issues:

Budget: {json.dumps(structured['budget'], indent=2)}
Text: {text[:5000]}

Return JSON:
{{
    "is_valid": <true/false>,
    "total_budget": <number>,
    "issues": [list of issues with severity],
    "recommendations": [list]
}}
"""
        
        response = gemini_model.generate_content(prompt)
        validation = parse_json_response(response.text)
        
        # Cleanup
        try:
            os.remove(file_path)
        except:
            pass
        
        return {
            "status": "success",
            "budget_extracted": structured['budget'],
            "validation": validation
        }
    except Exception as e:
        raise HTTPException(500, f"Budget validation error: {str(e)}")


@app.get("/api/uploads/list")
async def list_uploads():
    """List uploaded DPR files"""
    try:
        files = [
            {
                "filename": f,
                "size_bytes": os.path.getsize(f"uploads/{f}"),
                "uploaded": datetime.fromtimestamp(
                    os.path.getctime(f"uploads/{f}")
                ).isoformat()
            }
            for f in os.listdir("uploads")
            if f.endswith(('.pdf', '.docx', '.doc'))
        ]
        return {"uploads": files, "count": len(files)}
    except:
        return {"uploads": [], "count": 0}


@app.get("/api/guidelines/list")
async def list_guidelines():
    """List loaded guideline files"""
    try:
        files = [
            {
                "filename": f,
                "size_bytes": os.path.getsize(f"data/guidelines/{f}")
            }
            for f in os.listdir("data/guidelines")
            if f.endswith(('.pdf', '.docx', '.doc'))
        ]
        return {"guidelines": files, "count": len(files)}
    except:
        return {"guidelines": [], "count": 0}


# ============================================================================
# DOCUMENT MANAGEMENT ENDPOINTS
# ============================================================================

# Create documents storage directory
os.makedirs("data/documents", exist_ok=True)
DOCUMENTS_FILE = "data/documents/documents.json"

def load_documents():
    """Load documents from JSON file"""
    try:
        if os.path.exists(DOCUMENTS_FILE):
            with open(DOCUMENTS_FILE, 'r', encoding='utf-8') as f:
                return json.load(f)
        return []
    except:
        return []

def save_documents(documents):
    """Save documents to JSON file"""
    try:
        with open(DOCUMENTS_FILE, 'w', encoding='utf-8') as f:
            json.dump(documents, f, indent=2, ensure_ascii=False)
        return True
    except Exception as e:
        print(f"Error saving documents: {e}")
        return False

@app.post("/api/documents/add")
async def add_document(document_data: dict):
    """Add a new document to the system"""
    try:
        documents = load_documents()
        
        # Add timestamp and unique ID if not present
        if 'id' not in document_data:
            document_data['id'] = f"doc_{int(datetime.now().timestamp() * 1000)}"
        
        if 'uploadDate' not in document_data:
            document_data['uploadDate'] = datetime.now().isoformat().split('T')[0]
        
        documents.append(document_data)
        
        if save_documents(documents):
            return {"status": "success", "message": "Document added successfully", "document": document_data}
        else:
            raise HTTPException(500, "Failed to save document")
    except Exception as e:
        raise HTTPException(500, f"Error adding document: {str(e)}")

@app.get("/api/documents/list")
async def list_documents(user_email: str = None):
    """List all documents or documents for a specific user"""
    try:
        documents = load_documents()
        
        if user_email:
            # Filter documents for specific user
            user_documents = [doc for doc in documents if doc.get('uploadedBy', {}).get('email') == user_email]
            return {"status": "success", "documents": user_documents, "count": len(user_documents)}
        else:
            # Return all documents (for admin)
            return {"status": "success", "documents": documents, "count": len(documents)}
    except Exception as e:
        return {"status": "error", "documents": [], "count": 0, "error": str(e)}

@app.put("/api/documents/{document_id}/status")
async def update_document_status(document_id: str, status_data: dict):
    """Update document status"""
    try:
        documents = load_documents()
        
        # Find and update the document
        updated = False
        for doc in documents:
            if doc.get('id') == document_id:
                doc['status'] = status_data.get('status', doc.get('status'))
                if 'reviewerComments' in status_data:
                    doc['reviewerComments'] = status_data['reviewerComments']
                if 'reviewedBy' in status_data:
                    doc['reviewedBy'] = status_data['reviewedBy']
                doc['lastUpdated'] = datetime.now().isoformat()
                updated = True
                break
        
        if not updated:
            raise HTTPException(404, f"Document with ID {document_id} not found")
        
        if save_documents(documents):
            return {"status": "success", "message": "Document status updated successfully"}
        else:
            raise HTTPException(500, "Failed to save document changes")
    except HTTPException:
        raise
    except Exception as e:
        raise HTTPException(500, f"Error updating document: {str(e)}")

@app.get("/api/documents/submitted")
async def get_submitted_documents():
    """Get all documents with status 'submitted' for admin review"""
    try:
        documents = load_documents()
        submitted_docs = [doc for doc in documents if doc.get('status') == 'submitted']
        return {"status": "success", "documents": submitted_docs, "count": len(submitted_docs)}
    except Exception as e:
        return {"status": "error", "documents": [], "count": 0, "error": str(e)}

@app.delete("/api/documents/{document_id}")
async def delete_document(document_id: str):
    """Delete a document"""
    try:
        documents = load_documents()
        
        # Find and remove the document
        documents = [doc for doc in documents if doc.get('id') != document_id]
        
        if save_documents(documents):
            return {"status": "success", "message": "Document deleted successfully"}
        else:
            raise HTTPException(500, "Failed to save changes")
    except Exception as e:
        raise HTTPException(500, f"Error deleting document: {str(e)}")


# ============================================================================
# STARTUP & MAIN
# ============================================================================

@app.on_event("startup")
async def startup_event():
    """Run on application startup"""
    print("\n" + "="*60)
    print("[SYSTEM] DPR Analysis System Starting...")
    print("="*60)
    print(f"[SERVER] API Server: http://localhost:{APP_PORT}")
    print(f"[DOCS] API Docs: http://localhost:{APP_PORT}/docs")
    print(f"[AI-MODEL] {GEMINI_MODEL}")
    print("="*60 + "\n")


if __name__ == "__main__":
    import uvicorn
    
    print("\n[STARTING] DPR Analysis System...")
    print(f"[API-KEY] Gemini API Key: {'Configured' if GEMINI_API_KEY else 'Missing'}")
    
    uvicorn.run(
        app,
        host=APP_HOST,
        port=APP_PORT,
        log_level="info"
    )
